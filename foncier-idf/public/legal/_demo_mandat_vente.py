"""
Démo : génère un mandat de vente rempli avec des données fictives
pour visualiser le rendu final côté Samuel.

Reprend le pipeline backend (docxtemplater côté JS) mais en Python
via un simple string replacement sur tous les paragraphes.

Output : C:\\Users\\PC\\OneDrive\\Documents\\DEMO-mandat-vente-rempli.pdf
"""

from datetime import date
from pathlib import Path
from docx import Document

ROOT = Path(__file__).parent
TEMPLATE = ROOT / "templates" / "mandat-vente.template.docx"
OUTPUT_DOCX = Path(r"C:\Users\PC\OneDrive\Documents\DEMO-mandat-vente-rempli-v2.docx")
OUTPUT_PDF = Path(r"C:\Users\PC\OneDrive\Documents\DEMO-mandat-vente-rempli-v2.pdf")


# Forcer le mois en français (strftime("%B") rend en anglais si locale système EN)
MOIS_FR = {
    1: "janvier", 2: "février", 3: "mars", 4: "avril",
    5: "mai", 6: "juin", 7: "juillet", 8: "août",
    9: "septembre", 10: "octobre", 11: "novembre", 12: "décembre",
}


def fmt_date_fr(d: date) -> str:
    """Formate une date en français : '12 juin 2026'."""
    return f"{d.day} {MOIS_FR[d.month]} {d.year}"


# ============================================================
# Données fictives — vendeur HNWI à Paris 17e
# ============================================================
today = date.today()
# Ajout sûr de 3 mois (gère le passage d'année)
end_year = today.year + (today.month + 3 - 1) // 12
end_month = ((today.month + 3 - 1) % 12) + 1
# Cap day to last day of target month if needed
import calendar
end_day = min(today.day, calendar.monthrange(end_year, end_month)[1])
end_date = date(end_year, end_month, end_day)

TAGS = {
    # Header
    "numero_registre": "20260042",
    "date_mandat": fmt_date_fr(today),
    "date_fin": fmt_date_fr(end_date),

    # Mandataire — Diara
    "prenom_mandataire": "Diara",
    "nom_mandataire": "CAMARA",
    "numero_rsac_mandataire": "Paris RSAC 2026 B 12345",
    "numero_attestation": "ATT-EUR-2026-001",

    # Client vendeur
    "client_civilite": "M.",
    "client_nom": "DUPONT",
    "client_prenom": "Pierre-Antoine",
    "client_date_naissance": "15 mars 1972",
    "client_lieu_naissance": "Boulogne-Billancourt (92)",
    "client_nationalite": "Française",
    "client_adresse": "12 rue de Tocqueville, 75017 Paris",
    "client_telephone": "06 12 34 56 78",
    "client_email": "p.dupont@famille-dupont.fr",

    # Bien
    "bien_type": "Appartement Haussmannien T4",
    "bien_adresse": "8 avenue de Wagram, 75017 Paris (4e étage)",
    "bien_surface": "112",
    "bien_reference": "Cadastre Section AC parcelle 142, lot copro n°48",
    "bien_description": (
        "Appartement traversant Est/Ouest, 4 pièces principales, "
        "double salon, 2 chambres dont une suite parentale, "
        "cuisine équipée, salle de bains, dressing, balcon filant. "
        "Parquet point de Hongrie, moulures, cheminée d'origine. "
        "Cave et chambre de service au 6e."
    ),

    # Modalité
    "mandat_modalite": "exclusif",
    "duree_mois": "3",

    # Prix
    "prix_net_vendeur": "1 500 000",
    "prix_net_vendeur_lettres": "un million cinq cent mille euros",

    # Honoraires
    "commission_pct": "5",
    "commission_eur": "75 000",
    "commission_eur_lettres": "soixante-quinze mille euros",
}


def replace_in_paragraph(p, tags):
    """Remplace les tags {champ} dans un paragraphe."""
    full = "".join(r.text for r in p.runs)
    new = full
    for k, v in tags.items():
        new = new.replace("{" + k + "}", str(v))
    if new != full and p.runs:
        # Vide tous les runs sauf le premier, écrit le nouveau texte
        first = p.runs[0]
        for r in p.runs[1:]:
            r.text = ""
        first.text = new


def main():
    if not TEMPLATE.exists():
        raise SystemExit(f"Template introuvable : {TEMPLATE}")

    doc = Document(str(TEMPLATE))

    # Parcourir tous les paragraphes (corps + tables)
    for p in doc.paragraphs:
        replace_in_paragraph(p, TAGS)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p, TAGS)

    doc.save(str(OUTPUT_DOCX))
    print(f"OK -> {OUTPUT_DOCX}")

    # Conversion en PDF
    from docx2pdf import convert
    convert(str(OUTPUT_DOCX), str(OUTPUT_PDF))
    print(f"OK -> {OUTPUT_PDF}")


if __name__ == "__main__":
    main()
