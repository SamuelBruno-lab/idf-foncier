"""
Corrige l'article 12 du contrat final4.docx pour résoudre la
contradiction entre :
  - Article 12 actuel (PI = Mandant = Eurealimmo) ❌
  - Articles 8.2 bis / 16.6 / 16.6 bis / 16.6 ter (PI = DATAMERRY SAS) ✅

Le nouvel article 12 attribue la PI à DATAMERRY SAS et précise
la cascade de licence : DATAMERRY SAS → Eurealimmo SARL → Mandataire.

Input  : contrat-diara-final4.docx
Output : contrat-diara-final5.docx (cohérence PI restaurée)
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ONEDRIVE = Path(r"C:\Users\PC\OneDrive\Documents")
INPUT = ONEDRIVE / "contrat-diara-final4.docx"
OUTPUT = ONEDRIVE / "contrat-diara-final5.docx"

BLACK = RGBColor(0x00, 0x00, 0x00)

# Nouveau texte article 12 (3 paragraphes — cascade DATAMERRY → Mandant → Mandataire)
NEW_ART_12_P1 = (
    "La marque DATAMERRY® et ses dérivés, les développements logiciels "
    "associés, les bases de données agrégées (DVF, OSM, Wikidata, IDFM "
    "PRIM, INSEE, cadastre, permis de construire), les algorithmes de "
    "clusterisation HDBSCAN, les modèles CAPM-DCF, le moteur d'estimation "
    "immobilière, le module de notarisation cryptographique des mandats "
    "sur blockchain Solana ainsi que l'ensemble des éléments propriétaires "
    "de la plateforme demeurent la propriété exclusive et inaliénable de "
    "DATAMERRY SAS (société en cours de constitution à la date du présent "
    "contrat, cf. article 16.6 ter), entité juridique distincte du Mandant."
)

NEW_ART_12_P2 = (
    "Le Mandant (EUREALIMMO SARL) bénéficie, pour les besoins exclusifs "
    "de l'exploitation du réseau Eurealimmo Réseau et de l'activité "
    "d'agent immobilier titulaire de la carte T, d'une licence d'usage "
    "non exclusive, non transférable et révocable, consentie par "
    "DATAMERRY SAS dans les conditions de l'article 16.6 du présent "
    "contrat, sans préjudice du régime de subsidiarité prévu à l'article "
    "16.6 ter au profit de M. Samuel BRUNO à titre personnel."
)

NEW_ART_12_P3 = (
    "Le Mandataire ne dispose, sur lesdits éléments, que d'un droit "
    "d'utilisation strictement personnel, non exclusif, non cessible et "
    "limité à la durée du présent contrat, sous forme d'une "
    "sous-licence concédée par le Mandant au titre de son propre droit "
    "d'usage. Toute reproduction, copie, modification, rétro-ingénierie, "
    "extraction de bases de données ou tentative de contournement des "
    "protections techniques est strictement prohibée et constituera un "
    "manquement grave au présent contrat au sens de l'article 13.4."
)


def replace_paragraph_text(p, new_text):
    if not p.runs:
        p.add_run("")
    first_run = p.runs[0]
    for run in p.runs[1:]:
        run.text = ""
    first_run.text = new_text


def style_run(run, bold=False, italic=False, size=11, color=BLACK, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_paragraph_after(reference_paragraph, text, body=True):
    """Insère un paragraphe APRÈS le paragraphe de référence (manip XML)."""
    from copy import deepcopy
    # Clone le paragraphe de référence pour conserver le style global du doc
    new_p_xml = deepcopy(reference_paragraph._element)
    # Insère après
    reference_paragraph._element.addnext(new_p_xml)
    # Récupère le nouvel objet Paragraph
    from docx.text.paragraph import Paragraph
    new_p = Paragraph(new_p_xml, reference_paragraph._parent)
    # Vide le contenu cloné et remet le texte
    for r in list(new_p.runs):
        r.text = ""
    new_p.runs[0].text = text if new_p.runs else ""
    if not new_p.runs:
        new_p.add_run(text)
    # Style
    run = new_p.runs[0]
    style_run(run, size=11)
    if body:
        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        new_p.paragraph_format.space_after = Pt(6)
        new_p.paragraph_format.line_spacing = 1.25
    return new_p


def find_paragraph_by_contains(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


def main():
    if not INPUT.exists():
        raise SystemExit(f"Introuvable : {INPUT}")

    doc = Document(str(INPUT))

    # 1) Trouver le paragraphe actuel de l'art. 12 (contient "La marque DATAMERRY")
    p_12 = find_paragraph_by_contains(doc, "La marque DATAMERRY")
    if p_12 is None:
        raise SystemExit("Ancre article 12 non trouvee")
    # Remplacer son contenu par le para 1 NEW
    replace_paragraph_text(p_12, NEW_ART_12_P1)
    print("[OK] Article 12 para 1 remplace : PI = DATAMERRY SAS")

    # 2/3) Ajouter para 2 et para 3 APRES le para 1
    p_p2 = add_paragraph_after(p_12, NEW_ART_12_P2)
    print("[OK] Article 12 para 2 ajoute : licence DATAMERRY vers Mandant")
    add_paragraph_after(p_p2, NEW_ART_12_P3)
    print("[OK] Article 12 para 3 ajoute : sous-licence Mandant vers Mandataire")

    doc.save(str(OUTPUT))
    print(f"\nOK -> {OUTPUT}")


if __name__ == "__main__":
    main()
