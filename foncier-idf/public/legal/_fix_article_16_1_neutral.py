"""
Reformule l'article 16.1 du contrat pour rendre la blockchain
NEUTRE en techno (Bitcoin via OpenTimestamps OU Solana, au choix).

Permet de livrer la Phase 1 (OpenTimestamps Bitcoin, gratuit) sans
s'engager sur Solana spécifiquement.

Input  : contrat-diara-final5.docx
Output : contrat-diara-final6.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ONEDRIVE = Path(r"C:\Users\PC\OneDrive\Documents")
INPUT = ONEDRIVE / "contrat-diara-final5.docx"
OUTPUT = ONEDRIVE / "contrat-diara-final6.docx"

BLACK = RGBColor(0x00, 0x00, 0x00)

# Nouveau texte 16.1 — neutre vis-à-vis de la techno blockchain
NEW_16_1 = (
    "Le Mandant s'engage à mettre en place, sur une base mensuelle, "
    "un mécanisme d'ancrage cryptographique des empreintes SHA-256 des "
    "mandats et opérations détenus dans le système, sur un registre "
    "public et immuable de type blockchain. Cet ancrage vise à garantir "
    "la traçabilité, l'auditabilité et la non-répudiation des opérations "
    "sans porter atteinte à la confidentialité des données personnelles "
    "(article 16.2). Le choix du registre cryptographique est laissé à "
    "la libre appréciation du Mandant selon les évolutions technologiques "
    "et le contexte économique du moment. À titre indicatif, ce registre "
    "pourra prendre la forme : (i) en Phase 1, d'une notarisation via le "
    "protocole Open Timestamps ancré sur la blockchain Bitcoin (standard "
    "ouvert, gratuit, immédiatement déployable) ; (ii) en Phase 2, d'un "
    "smart contract dédié sur la blockchain Solana ou tout autre réseau "
    "Layer 1 jugé pertinent par le Mandant. La transition entre les "
    "phases ne saurait emporter modification des autres stipulations du "
    "présent contrat."
)


def replace_paragraph_text(p, new_text):
    if not p.runs:
        p.add_run("")
    first_run = p.runs[0]
    for run in p.runs[1:]:
        run.text = ""
    first_run.text = new_text


def find_paragraph_by_contains(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


def main():
    if not INPUT.exists():
        raise SystemExit(f"Introuvable : {INPUT}")

    doc = Document(str(INPUT))

    # Trouver le paragraphe 16.1 (contenu actuel)
    p_161 = find_paragraph_by_contains(doc, "Merkle Root mensuel")
    if p_161 is None:
        p_161 = find_paragraph_by_contains(doc, "16.1")
    if p_161 is None:
        raise SystemExit("Ancre 16.1 non trouvee")
    replace_paragraph_text(p_161, NEW_16_1)
    print("[OK] Article 16.1 reformule : techno-neutre (Bitcoin Phase 1 / Solana Phase 2)")

    doc.save(str(OUTPUT))
    print(f"\nOK -> {OUTPUT}")


if __name__ == "__main__":
    main()
