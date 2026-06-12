"""
Précise dans la clause 1.bis du contrat Diara que la carte T autorise
LES DEUX SENS de la location :
  (a) mise en location (recherche locataire pour bailleur)
  (b) recherche de bien locatif (pour le compte du candidat locataire)

Input  : contrat-diara-final8.docx
Output : contrat-diara-final9.docx
"""

from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

ONEDRIVE = Path(r"C:\Users\PC\OneDrive\Documents")
INPUT = ONEDRIVE / "contrat-diara-final8.docx"
OUTPUT = ONEDRIVE / "contrat-diara-final9.docx"

BLACK = RGBColor(0x00, 0x00, 0x00)


# Remplace le paragraphe "A l'inverse..." par 4 paragraphes structures
NEW_PARA_HEADER = (
    "À l'inverse, demeurent autorisées et entrent dans le périmètre de "
    "la carte T les opérations suivantes, qui relèvent toutes de "
    "l'entremise immobilière au sens de l'article 1er de la loi Hoguet :"
)

NEW_BULLETS = [
    "(a) la mise en location, entendue comme l'entremise initiale visant la recherche d'un LOCATAIRE pour le compte d'un BAILLEUR ;",
    "(b) la recherche de bien immobilier à louer (résidentiel, commercial ou professionnel) pour le compte d'un candidat LOCATAIRE ou d'une entreprise preneuse.",
]

NEW_PARA_FOOTER = (
    "Les opérations visées aux (a) et (b) ci-dessus s'achèvent à la "
    "signature du bail exclusivement, à l'exclusion de toute opération "
    "d'administration ou de gestion subséquente (encaissement de loyers, "
    "charges, états des lieux post-bail, etc.), laquelle relèverait de "
    "la carte G dont le Mandant n'est pas titulaire."
)


def style_run(run, bold=False, italic=False, size=11, color=BLACK, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def replace_paragraph_text(p, new_text):
    if not p.runs:
        p.add_run("")
    first_run = p.runs[0]
    for r in p.runs[1:]:
        r.text = ""
    first_run.text = new_text


def add_paragraph_after(reference_paragraph, text, *, bold=False, bullet=False):
    """Insère un nouveau paragraphe APRÈS le paragraphe de référence."""
    new_p_xml = deepcopy(reference_paragraph._element)
    reference_paragraph._element.addnext(new_p_xml)
    new_p = Paragraph(new_p_xml, reference_paragraph._parent)
    # Vide le contenu cloné
    for r in list(new_p.runs):
        r.text = ""
    if new_p.runs:
        new_p.runs[0].text = text
    else:
        new_p.add_run(text)

    run = new_p.runs[0]
    if bullet:
        new_p.style = "List Bullet"
        style_run(run, size=11)
        new_p.paragraph_format.space_after = Pt(3)
        new_p.paragraph_format.line_spacing = 1.2
    else:
        style_run(run, bold=bold, size=11)
        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        new_p.paragraph_format.space_after = Pt(6)
        new_p.paragraph_format.line_spacing = 1.25
    return new_p


def find_paragraph_by_contains(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


def main():
    if not INPUT.exists():
        raise SystemExit(f"Introuvable : {INPUT}")

    doc = Document(str(INPUT))

    # Trouve le paragraphe "À l'inverse, demeure autorisée..."
    p_inverse = find_paragraph_by_contains(doc, "demeure autorisée")
    if p_inverse is None:
        p_inverse = find_paragraph_by_contains(doc, "demeure autoris")
    if p_inverse is None:
        raise SystemExit("Ancre 'à l'inverse...' non trouvée")

    # 1. Remplace le contenu par l'introduction
    replace_paragraph_text(p_inverse, NEW_PARA_HEADER)
    print("[OK] Paragraphe d'introduction reformulé")

    # 2. Insère les 2 bullets après
    last = p_inverse
    for bullet_text in NEW_BULLETS:
        last = add_paragraph_after(last, bullet_text, bullet=True)
    print("[OK] 2 bullets (a) et (b) inseres")

    # 3. Insère le paragraphe de conclusion
    add_paragraph_after(last, NEW_PARA_FOOTER)
    print("[OK] Paragraphe de conclusion ajoute")

    doc.save(str(OUTPUT))
    print(f"\nOK -> {OUTPUT}")


if __name__ == "__main__":
    main()
