"""
Injecte l'ARTICLE 11 BIS — PROPRIÉTÉ DES DONNÉES, MARQUES ET PROSPECTS
dans le contrat V2 validé par Diara, sans toucher au reste.

Wording approuvé par WhatsApp Diara :
  "Ton site collabimo.com et tes prospects pré-mandat restent à toi.
   Les mandats que tu signes sous ma carte T appartiennent juridiquement
   à Eurealimmo (loi Hoguet l'exige), mais en pratique tu y as accès,
   95 % de commission. Fin de contrat = leads pré-mandat + marque Collabimo te suivent."

Input  : contrat-mandat-diara-eurealimmo V2.docx
Output : contrat-mandat-diara-eurealimmo V3.docx (+ .pdf)

Insertion : juste AVANT "ARTICLE 12 — PROPRIÉTÉ INTELLECTUELLE DATAMERRY"
"""

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ONEDRIVE = Path(r"C:\Users\PC\OneDrive\Documents")
INPUT = ONEDRIVE / "contrat-mandat-diara-eurealimmo V2.docx"
OUTPUT_DOCX = ONEDRIVE / "contrat-mandat-diara-eurealimmo V3.docx"
OUTPUT_PDF = ONEDRIVE / "contrat-mandat-diara-eurealimmo V3.pdf"

DM_DARK = RGBColor(0x06, 0x4E, 0x3B)
BLACK = RGBColor(0x00, 0x00, 0x00)


def style_run(run, *, bold=False, italic=False, size=11, color=BLACK, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def insert_heading_1(reference_paragraph, text):
    p = reference_paragraph.insert_paragraph_before("")
    run = p.add_run(text)
    style_run(run, bold=True, size=13, color=DM_DARK)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    return p


def insert_heading_2(reference_paragraph, text):
    p = reference_paragraph.insert_paragraph_before("")
    run = p.add_run(text)
    style_run(run, bold=True, size=11, color=BLACK)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    return p


def insert_body(reference_paragraph, text, *, italic=False):
    p = reference_paragraph.insert_paragraph_before("")
    run = p.add_run(text)
    style_run(run, italic=italic, size=11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    return p


def insert_bullet(reference_paragraph, text):
    """Insère un alinéa indenté style (i), (ii), (iii)."""
    p = reference_paragraph.insert_paragraph_before("")
    run = p.add_run(text)
    style_run(run, size=11)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    return p


def find_anchor(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().upper().startswith(prefix.upper()):
            return p
    raise SystemExit(f"Ancre introuvable : {prefix!r}")


def main():
    if not INPUT.exists():
        raise SystemExit(f"Fichier source introuvable : {INPUT}")

    doc = Document(str(INPUT))

    # Ancre : juste avant ARTICLE 12 (= insertion entre article 11 et article 12)
    anchor = find_anchor(doc, "ARTICLE 12")
    print(f"OK -> ancre trouvée : {anchor.text!r}")

    # ===========================================================
    # ARTICLE 11 BIS
    # ===========================================================
    insert_heading_1(
        anchor,
        "ARTICLE 11 BIS — PROPRIÉTÉ DES DONNÉES, MARQUES ET PROSPECTS (clarification interprétative)",
    )

    # --- 11 bis.1 ---
    insert_heading_2(anchor, "11 bis.1 — Site collabimo.com et Prospects Pré-Mandat")
    insert_body(
        anchor,
        "Le site internet collabimo.com et l'ensemble des prospects, "
        "contacts et données commerciales (ci-après désignés les "
        "« Prospects Pré-Mandat ») collectés par le Mandataire via la marque "
        "commerciale COLLABIMO préalablement à toute signature d'un mandat "
        "de transaction immobilière avec un client vendeur, demeurent la "
        "propriété exclusive du Mandataire en sa qualité de représentante "
        "légale de Collabimo SARL. Ces éléments ne constituent pas des "
        "biens du Mandant et ne sont pas couverts par l'obligation de "
        "non-sollicitation prévue à l'article 9.1.",
    )

    # --- 11 bis.2 ---
    insert_heading_2(
        anchor, "11 bis.2 — Mandats signés sous carte T Eurealimmo"
    )
    insert_body(
        anchor,
        "Conformément aux obligations résultant de la loi n° 70-9 du "
        "2 janvier 1970 (dite « loi Hoguet ») et du décret n° 72-678 du "
        "20 juillet 1972, les mandats de transaction immobilière conclus "
        "par le Mandataire avec des clients vendeurs au cours de "
        "l'exécution du présent contrat sont juridiquement la propriété du "
        "Mandant, EUREALIMMO SARL, titulaire de la carte professionnelle T "
        "n° CPI 7501 2024 000 219. Ces mandats sont inscrits au registre "
        "des mandats tenu par le Mandant.",
    )
    insert_body(anchor, "Le Mandataire conserve néanmoins :")
    insert_bullet(
        anchor,
        "(i) un accès opérationnel intégral à ces mandats via la plateforme DATAMERRY® ;",
    )
    insert_bullet(
        anchor,
        "(ii) la rétrocession de quatre-vingt-quinze pour cent (95 %) des "
        "commissions encaissées par le Mandant sur les mandats qu'elle a "
        "apportés, conformément à l'article 6.1 ;",
    )
    insert_bullet(
        anchor,
        "(iii) le droit de représentation effective et exclusive auprès du "
        "client vendeur correspondant.",
    )

    # --- 11 bis.3 ---
    insert_heading_2(anchor, "11 bis.3 — Sortie du contrat")
    insert_body(
        anchor,
        "À l'issue du présent contrat, quelle qu'en soit la cause :",
    )
    insert_bullet(
        anchor,
        "(i) Le Mandataire reprend la pleine et libre disposition de la "
        "marque COLLABIMO, du site collabimo.com, et de l'ensemble des "
        "Prospects Pré-Mandat au sens de l'article 11 bis.1 ;",
    )
    insert_bullet(
        anchor,
        "(ii) Les mandats actifs au jour de la fin du contrat demeurent "
        "attachés à la carte T du Mandant et continuent d'être exécutés "
        "selon leurs termes propres jusqu'à leur extinction. Le Mandataire "
        "conserve la rétrocession à quatre-vingt-quinze pour cent (95 %) "
        "sur les commissions perçues par le Mandant au titre de ces "
        "mandats actifs ;",
    )
    insert_bullet(
        anchor,
        "(iii) Le Mandataire reste tenu de l'obligation de non-sollicitation "
        "prévue à l'article 9.1, portant exclusivement sur les clients "
        "vendeurs effectivement sous mandat actif Eurealimmo au jour de la "
        "fin du contrat, pour une durée de douze (12) mois calendaires. "
        "Les Prospects Pré-Mandat ne sont pas concernés par cette obligation.",
    )

    # --- 11 bis.4 ---
    insert_heading_2(anchor, "11 bis.4 — Caractère interprétatif")
    insert_body(
        anchor,
        "Les présentes stipulations sont déclaratives et n'ont pas pour "
        "objet de créer de nouveaux droits ou obligations, mais de "
        "clarifier l'articulation des règles déjà définies aux articles 6, "
        "9 et 11 du présent contrat à la lumière de la double qualité du "
        "Mandataire, représentante légale de Collabimo SARL et associée "
        "fondatrice n° 1 d'EUREALIMMO RÉSEAU. En cas de contradiction "
        "entre le présent article et un autre article du présent contrat, "
        "le présent article prévaut.",
    )

    doc.save(str(OUTPUT_DOCX))
    print(f"OK -> {OUTPUT_DOCX}")

    # Conversion PDF (best-effort)
    try:
        from docx2pdf import convert
        convert(str(OUTPUT_DOCX), str(OUTPUT_PDF))
        print(f"OK -> {OUTPUT_PDF}")
    except Exception as e:
        print(f"PDF conversion ignorée : {e}")


if __name__ == "__main__":
    main()
