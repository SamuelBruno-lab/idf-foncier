"""
Ajoute à l'article 1 du contrat Diara une exclusion expresse des
opérations relevant de la carte professionnelle G (Gestion immobilière).

Périmètre Eurealimmo = carte T UNIQUEMENT.

Input  : contrat-diara-final7.docx
Output : contrat-diara-final8.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ONEDRIVE = Path(r"C:\Users\PC\OneDrive\Documents")
INPUT = ONEDRIVE / "contrat-diara-final7.docx"
OUTPUT = ONEDRIVE / "contrat-diara-final8.docx"

BLACK = RGBColor(0x00, 0x00, 0x00)
DM_DARK = RGBColor(0x06, 0x4E, 0x3B)


# ===========================================================
# Nouveau paragraphe 1.bis — Exclusion carte G
# ===========================================================

CLAUSE_TITLE = "1.bis — Exclusion expresse des opérations relevant de la carte G (Gestion immobilière)"

CLAUSE_BODY_1 = (
    "Sont expressément exclues du périmètre du présent contrat toutes "
    "opérations relevant de la carte professionnelle G (« Gestion "
    "immobilière ») au sens de l'article 1er de la loi Hoguet "
    "n° 70-9 du 2 janvier 1970, le Mandant n'étant titulaire QUE de "
    "la carte professionnelle T (Transactions) n° CPI 7501 2024 000 219. "
    "Sont notamment exclus du présent mandat :"
)

CLAUSE_BULLETS = [
    "(i) l'administration de biens d'autrui ;",
    "(ii) la gestion locative, en ce compris l'encaissement de loyers, de charges, les régularisations annuelles, les états des lieux post-bail et le traitement des congés ;",
    "(iii) la fonction de syndic de copropriété ;",
    "(iv) tout maniement de fonds clients au sens de l'article 3-2 de la loi Hoguet.",
]

CLAUSE_BODY_2 = (
    "Le Mandataire s'interdit expressément de proposer, négocier ou "
    "signer de tels mandats au nom du Mandant. La méconnaissance de "
    "cette interdiction constitue une faute grave au sens de "
    "l'article 13 alinéa 2 du présent contrat et de l'article L134-13 "
    "du Code de commerce."
)

CLAUSE_BODY_3 = (
    "À l'inverse, demeure autorisée et entre dans le périmètre de la "
    "carte T la mise en location, entendue comme l'entremise initiale "
    "visant la recherche d'un locataire pour le compte d'un bailleur, "
    "jusqu'à la signature du bail exclusivement, à l'exclusion de "
    "toute opération de gestion subséquente."
)


# ===========================================================
# Utils
# ===========================================================

def style_run(run, bold=False, italic=False, size=11, color=BLACK, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_paragraph_before(reference_paragraph, text, *,
                         bold=False, bullet=False, size=11):
    new_p = reference_paragraph.insert_paragraph_before("")
    if bullet:
        new_p.style = "List Bullet"
        run = new_p.add_run(text)
        style_run(run, size=size)
        new_p.paragraph_format.space_after = Pt(3)
        new_p.paragraph_format.line_spacing = 1.2
    else:
        run = new_p.add_run(text)
        style_run(run, bold=bold, size=size)
        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        new_p.paragraph_format.space_after = Pt(6)
        new_p.paragraph_format.line_spacing = 1.25
    return new_p


def find_paragraph_starting_with(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


# ===========================================================
# Main
# ===========================================================

def main():
    if not INPUT.exists():
        raise SystemExit(f"Introuvable : {INPUT}")

    doc = Document(str(INPUT))

    # Insertion juste avant "ARTICLE 2"
    p_2 = find_paragraph_starting_with(doc, "ARTICLE 2")
    if p_2 is None:
        raise SystemExit("Ancre ARTICLE 2 non trouvee")

    add_paragraph_before(p_2, CLAUSE_TITLE, bold=True)
    add_paragraph_before(p_2, CLAUSE_BODY_1)
    for b in CLAUSE_BULLETS:
        add_paragraph_before(p_2, b, bullet=True)
    add_paragraph_before(p_2, CLAUSE_BODY_2)
    add_paragraph_before(p_2, CLAUSE_BODY_3)

    print("[OK] Insertion clause 1.bis : exclusion carte G")

    doc.save(str(OUTPUT))
    print(f"OK -> {OUTPUT}")


if __name__ == "__main__":
    main()
