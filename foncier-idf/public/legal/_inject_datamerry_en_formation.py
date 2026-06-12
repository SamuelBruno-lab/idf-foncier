"""
Ajoute au contrat final3.docx les précisions :
  - DATAMERRY SAS est "en cours de constitution" à la date du contrat
  - Clause de subsidiarité : si non immatriculée au RCS dans les 4 mois,
    Samuel BRUNO personne physique reprend les droits

Input  : contrat-diara-final3.docx
Output : contrat-diara-final4.docx (version définitive à envoyer)
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ONEDRIVE = Path(r"C:\Users\PC\OneDrive\Documents")
INPUT = ONEDRIVE / "contrat-diara-final3.docx"
OUTPUT = ONEDRIVE / "contrat-diara-final4.docx"

BLACK = RGBColor(0x00, 0x00, 0x00)


# ---------------------------------------------------------------------------
# Contenu à insérer / remplacer
# ---------------------------------------------------------------------------

# Nouveau 8.2 bis : ajoute "(en cours de constitution)" + mention subsidiarité
NEW_8_2_BIS = (
    "Pour lever toute ambiguïté, la cession isolée des parts ou du fonds de "
    "commerce de DATAMERRY SAS (société en cours de constitution à la date "
    "du présent contrat — cf. article 16.6 ter), entité juridique distincte "
    "d'EUREALIMMO SARL et titulaire de la plateforme technologique "
    "comprenant notamment : (i) le moteur d'estimation immobilière, "
    "(ii) les bases de données agrégées (DVF, OSM, Wikidata, IDFM PRIM, "
    "INSEE, cadastre, permis de construire), (iii) le module de notarisation "
    "cryptographique des mandats sur blockchain Solana (algorithmes, smart "
    "contracts, infrastructure on-chain, interfaces de génération des "
    "certificats), et (iv) la marque DATAMERRY® et l'ensemble de ses "
    "dérivés, ne constitue PAS un fait générateur de la prime de cession "
    "prévue au présent article."
)

# Nouveau 16.6 : ajoute "(en cours de constitution)" mention
NEW_16_6 = (
    "Le module de notarisation blockchain est et restera entièrement GRATUIT "
    "pour le Mandataire pendant toute la durée du présent contrat, y compris "
    "en cas de renouvellement. Le module, les smart contracts associés, leur "
    "code source, leur infrastructure on-chain et leurs développements "
    "ultérieurs demeurent la propriété exclusive et inaliénable de "
    "DATAMERRY SAS (société en cours de constitution à la date du présent "
    "contrat, dont les conditions de substitution éventuelle sont précisées "
    "à l'article 16.6 ter), distincte juridiquement d'EUREALIMMO SARL. "
    "EUREALIMMO SARL en bénéficie au titre d'une licence d'usage non "
    "exclusive, non transférable et révocable consentie par DATAMERRY SAS "
    "pour les besoins exclusifs de l'exploitation du réseau Eurealimmo "
    "Réseau. DATAMERRY SAS se réserve la faculté de valoriser, céder ou "
    "exploiter commercialement ledit module auprès de tiers (notamment "
    "d'autres cabinets immobiliers, via le Solana Mobile dApp Store ou par "
    "cession à un acquéreur tiers), sans que cela puisse être assimilé à "
    "une cession au sens des articles 8 et 8.2 bis du présent contrat."
)

# Nouvelle clause 16.6 ter (subsidiarité société en formation)
NEW_16_6_TER_TITLE = "16.6 ter — Société en cours de constitution et clause de subsidiarité"
NEW_16_6_TER_BODY_1 = (
    "Les Parties reconnaissent expressément que, à la date de signature du "
    "présent contrat, DATAMERRY SAS est en cours de constitution et n'est "
    "pas encore immatriculée au Registre du Commerce et des Sociétés. Les "
    "stipulations du présent contrat faisant référence à DATAMERRY SAS sont "
    "consenties dans cette perspective, M. Samuel BRUNO agissant en qualité "
    "de fondateur et représentant légal en formation."
)
NEW_16_6_TER_BODY_2 = (
    "Dans l'hypothèse où DATAMERRY SAS ne serait pas effectivement "
    "immatriculée au Registre du Commerce et des Sociétés dans un délai de "
    "QUATRE (4) MOIS à compter de la date d'effet du présent contrat, la "
    "propriété intellectuelle relative au module de notarisation blockchain "
    "et à l'ensemble de ses composants (algorithmes, smart contracts, "
    "infrastructure on-chain, code source, bases de données associées et, "
    "le cas échéant, marque DATAMERRY®) sera réputée détenue à titre "
    "personnel par M. Samuel BRUNO, qui assumera dès lors, en lieu et place "
    "de DATAMERRY SAS, l'intégralité des droits et obligations afférents au "
    "présent contrat et à la licence d'usage consentie à EUREALIMMO SARL."
)
NEW_16_6_TER_BODY_3 = (
    "M. Samuel BRUNO conserve, dans cette hypothèse, la faculté de "
    "transférer ultérieurement ladite propriété intellectuelle à toute "
    "entité juridique de son choix (par voie d'apport en nature, de cession "
    "ou de tout autre mécanisme légal), sans que cette substitution "
    "n'affecte les droits et obligations du Mandataire au titre du présent "
    "contrat. Le Mandataire en sera informé par écrit dans un délai de "
    "trente (30) jours suivant l'opération."
)
NEW_16_6_TER_BODY_4 = (
    "Cette clause de subsidiarité ne modifie en aucune manière les "
    "stipulations des articles 8 (clause de cession) et 8.2 bis (exclusion "
    "DATAMERRY) du présent contrat, l'identité juridique du titulaire de la "
    "plateforme technologique étant sans incidence sur le périmètre des "
    "droits du Mandataire."
)


# ---------------------------------------------------------------------------
# Utilitaires
# ---------------------------------------------------------------------------

def replace_paragraph_text(p, new_text):
    if not p.runs:
        p.add_run("")
    first_run = p.runs[0]
    for run in p.runs[1:]:
        run.text = ""
    first_run.text = new_text


def style_run(run, bold=False, italic=False, size=11, color=BLACK, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_paragraph_before(reference_paragraph, text, *, bold=False, body=False):
    new_p = reference_paragraph.insert_paragraph_before("")
    run = new_p.add_run(text)
    style_run(run, bold=bold, size=11)
    if body:
        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        new_p.paragraph_format.space_after = Pt(6)
        new_p.paragraph_format.line_spacing = 1.25
    return new_p


def find_paragraph_by_contains(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if not INPUT.exists():
        raise SystemExit(f"Introuvable : {INPUT}")

    doc = Document(str(INPUT))

    # 1) Remplacer le corps de 8.2 bis (avec mention "société en cours de constitution")
    p_82bis = find_paragraph_by_contains(doc, "Pour lever toute ambiguïté")
    if p_82bis is None:
        p_82bis = find_paragraph_by_contains(doc, "Pour lever toute ambigu")
    if p_82bis is None:
        raise SystemExit("Ancre 8.2 bis non trouvée")
    replace_paragraph_text(p_82bis, NEW_8_2_BIS)
    print("[OK] 8.2 bis : ajout 'société en cours de constitution + renvoi 16.6 ter'")

    # 2) Remplacer le corps de 16.6 (avec mention société en cours)
    p_166 = find_paragraph_by_contains(doc, "Le module de notarisation blockchain est et restera")
    if p_166 is None:
        raise SystemExit("Ancre 16.6 non trouvée")
    replace_paragraph_text(p_166, NEW_16_6)
    print("[OK] 16.6 : ajout 'société en cours de constitution + renvoi 16.6 ter'")

    # 3) Insérer 16.6 ter juste après 16.6 bis (avant 16.7)
    p_167 = find_paragraph_by_contains(doc, "16.7 — Faculté d'expérimentation")
    if p_167 is None:
        p_167 = find_paragraph_by_contains(doc, "16.7")
    if p_167 is None:
        raise SystemExit("Ancre 16.7 non trouvée")
    add_paragraph_before(p_167, NEW_16_6_TER_TITLE, bold=True)
    add_paragraph_before(p_167, NEW_16_6_TER_BODY_1, body=True)
    add_paragraph_before(p_167, NEW_16_6_TER_BODY_2, body=True)
    add_paragraph_before(p_167, NEW_16_6_TER_BODY_3, body=True)
    add_paragraph_before(p_167, NEW_16_6_TER_BODY_4, body=True)
    print("[OK] 16.6 ter : nouvelle clause société en formation + subsidiarité 4 mois")

    doc.save(str(OUTPUT))
    print(f"\nOK -> {OUTPUT}")


if __name__ == "__main__":
    main()
