"""
Ajoute au contrat final6.docx une nouvelle clause :
  Article 11 ter — Système de routing des leads (lead matching)

Cette clause inscrit noir sur blanc que :
  - Le Mandataire contrôle EXCLUSIVEMENT les paramètres de matching
  - Aucun lead n'est attribué à un tiers sans validation Diara
  - Audit trail complet (lead_match_history)
  - Cohérence avec art. 11 bis (propriété data) et 16.x (anonymisation)

Input  : contrat-diara-final6.docx
Output : contrat-diara-final7.docx
"""

from pathlib import Path
from copy import deepcopy

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph

ONEDRIVE = Path(r"C:\Users\PC\OneDrive\Documents")
INPUT = ONEDRIVE / "contrat-diara-final6.docx"
OUTPUT = ONEDRIVE / "contrat-diara-final7.docx"

DM_DARK = RGBColor(0x06, 0x4E, 0x3B)
BLACK = RGBColor(0x00, 0x00, 0x00)


# ===========================================================
# Contenu Article 11 ter
# ===========================================================

CLAUSE_TITLE = (
    "ARTICLE 11 TER — SYSTÈME DE ROUTING DES LEADS (LEAD MATCHING) "
    "ET CONTRÔLE EXCLUSIF DU MANDATAIRE"
)

P1_TITLE = "11 ter.1 — Principe du système de matching"
P1_BODY = (
    "Le Mandant met à disposition du Mandataire, via la plateforme "
    "DATAMERRY®, un module de matching géo-spatial permettant, pour "
    "chaque lead capturé sur les canaux propres du Mandataire "
    "(notamment collabimo.com), d'identifier automatiquement les "
    "membres du réseau Collabimo (vendeurs, acheteurs, mandataires, "
    "apporteurs d'affaires) géographiquement proches du bien concerné "
    "selon des paramètres définis par le Mandataire."
)

P2_TITLE = "11 ter.2 — Contrôle exclusif du Mandataire sur les paramètres"
P2_BODY = (
    "Le Mandataire dispose d'un contrôle EXCLUSIF sur l'intégralité "
    "des paramètres du module de matching, accessibles via son "
    "interface d'administration (notamment l'URL "
    "/cabinets/collabimo/admin/matching). Sont notamment paramétrables "
    "par le seul Mandataire :"
)

P2_BULLETS = [
    "(i) les rayons géographiques de recherche par type de membre (mandataires, apporteurs d'affaires, vendeurs, acheteurs) ;",
    "(ii) les plafonds de rayon par densité de zone (Paris intra-muros, grandes métropoles, autres zones) ;",
    "(iii) l'activation et les paramètres du mode adaptatif (élargissement progressif du rayon jusqu'à atteindre un nombre cible de matches) ;",
    "(iv) les limites d'affichage (nombre minimum et maximum de matches présentés par lead) ;",
    "(v) le message d'absence de match présenté à l'utilisateur final.",
]

P3_TITLE = "11 ter.3 — Arbitrage exclusif du Mandataire sur l'attribution"
P3_BODY = (
    "Aucun lead capturé via le système ne peut être attribué à un "
    "tiers (qu'il s'agisse d'un membre du réseau Collabimo ou d'un "
    "mandataire Eurealimmo Réseau autre que le Mandataire) sans "
    "validation expresse et préalable du Mandataire. Pour chaque "
    "lead, le Mandataire dispose des options suivantes : "
    "(i) conserver le lead pour son propre compte en vue de la "
    "signature d'un mandat sous EUREALIMMO conformément à l'article "
    "11 bis.2 ; (ii) attribuer le lead à un membre du réseau "
    "Collabimo identifié par le système ; (iii) procéder à un "
    "matching croisé (vendeur ↔ acheteur, par exemple) ; (iv) "
    "écarter le lead s'il est jugé non qualifié."
)

P4_TITLE = "11 ter.4 — Traçabilité et auditabilité"
P4_BODY = (
    "Chaque arbitrage du Mandataire est enregistré dans une table "
    "d'historique dédiée (lead_match_history), comportant a minima : "
    "l'identifiant du lead, le type de lead (vendeur ou acheteur), "
    "le rayon utilisé, le nombre de matches proposés, la décision "
    "prise, l'identité du membre attributaire le cas échéant, "
    "l'identité de l'utilisateur ayant procédé à l'arbitrage et "
    "l'horodatage. Le Mandataire peut, à tout moment, exporter "
    "cet historique au format CSV ou JSON via son interface "
    "d'administration."
)

P5_TITLE = "11 ter.5 — Confidentialité des données des membres tiers"
P5_BODY = (
    "Lors de l'affichage des matches, le système ne révèle que les "
    "informations strictement nécessaires à la décision du "
    "Mandataire : prénom, nom, type de membre, ville, code postal, "
    "spécialité, distance approximative et performance moyenne. "
    "Les coordonnées personnelles complètes (email, téléphone) ne "
    "sont divulguées qu'à compter de l'arbitrage explicite du "
    "Mandataire, dans le strict respect de l'article 11 bis et du "
    "RGPD article 5-1-c) (principe de minimisation)."
)

P6_TITLE = "11 ter.6 — Engagement de non-fuite par le Mandant"
P6_BODY = (
    "Le Mandant s'interdit expressément, pendant toute la durée du "
    "présent contrat et pendant les douze (12) mois suivant sa "
    "cessation, de proposer, transférer, vendre ou exploiter les "
    "leads capturés via les canaux propres du Mandataire à un tiers "
    "concurrent ou à un autre mandataire Eurealimmo Réseau, en "
    "dehors de la stricte exécution des arbitrages validés par le "
    "Mandataire au titre du présent article 11 ter."
)

P7_TITLE = "11 ter.7 — Évolutions du système"
P7_BODY = (
    "Le Mandant se réserve la faculté d'améliorer le système de "
    "matching (algorithmes, ergonomie de l'interface, ajout de "
    "filtres ou de critères) à des fins d'amélioration du service, "
    "sous réserve de ne pas réduire le périmètre de contrôle "
    "exclusif du Mandataire défini au présent article. Toute "
    "évolution substantielle est préalablement notifiée au "
    "Mandataire avec un délai raisonnable de prise en compte."
)


# ===========================================================
# Utilitaires
# ===========================================================

def style_run(run, bold=False, italic=False, size=11, color=BLACK, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_paragraph_before(reference_paragraph, text, *,
                         heading=False,
                         bold=False,
                         bullet=False,
                         size=11):
    new_p = reference_paragraph.insert_paragraph_before("")
    if heading:
        run = new_p.add_run(text)
        style_run(run, bold=True, size=13, color=DM_DARK)
        new_p.paragraph_format.space_before = Pt(18)
        new_p.paragraph_format.space_after = Pt(6)
        new_p.paragraph_format.keep_with_next = True
    elif bullet:
        new_p.style = "List Bullet"
        run = new_p.add_run(text)
        style_run(run, size=size)
        new_p.paragraph_format.space_after = Pt(3)
        new_p.paragraph_format.line_spacing = 1.2
    else:
        run = new_p.add_run(text)
        style_run(run, bold=bold, size=size)
        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        new_p.paragraph_format.space_after = Pt(6)
        new_p.paragraph_format.line_spacing = 1.25
    return new_p


def find_paragraph_starting_with(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


# ===========================================================
# Main
# ===========================================================

def main():
    if not INPUT.exists():
        raise SystemExit(f"Introuvable : {INPUT}")

    doc = Document(str(INPUT))

    # On insere apres l'article 11 bis : trouver l'article 12 et inserer avant
    # (qui suit immediatement 11 bis dans le contrat)
    p_12 = find_paragraph_starting_with(doc, "ARTICLE 12")
    if p_12 is None:
        raise SystemExit("Ancre ARTICLE 12 non trouvee")

    # Insertion en cascade avant p_12 (donc apres 11 bis)
    add_paragraph_before(p_12, CLAUSE_TITLE, heading=True)

    add_paragraph_before(p_12, P1_TITLE, bold=True)
    add_paragraph_before(p_12, P1_BODY)

    add_paragraph_before(p_12, P2_TITLE, bold=True)
    add_paragraph_before(p_12, P2_BODY)
    for b in P2_BULLETS:
        add_paragraph_before(p_12, b, bullet=True)

    add_paragraph_before(p_12, P3_TITLE, bold=True)
    add_paragraph_before(p_12, P3_BODY)

    add_paragraph_before(p_12, P4_TITLE, bold=True)
    add_paragraph_before(p_12, P4_BODY)

    add_paragraph_before(p_12, P5_TITLE, bold=True)
    add_paragraph_before(p_12, P5_BODY)

    add_paragraph_before(p_12, P6_TITLE, bold=True)
    add_paragraph_before(p_12, P6_BODY)

    add_paragraph_before(p_12, P7_TITLE, bold=True)
    add_paragraph_before(p_12, P7_BODY)

    print("[OK] Article 11 ter insere (8 sous-sections)")

    doc.save(str(OUTPUT))
    print(f"\nOK -> {OUTPUT}")


if __name__ == "__main__":
    main()
