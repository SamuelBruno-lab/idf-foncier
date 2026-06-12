"""
Injecte le Pack PROTECTION (art. 6.1 précisé + 13 bis + 13 ter + 13 quater)
dans le DOCX FINAL annoté par Samuel, sans toucher au reste.

Input  : contrat-diara-final1.docx (version Samuel)
Output : contrat-diara-final2.docx (version + Pack PROTECTION)

Stratégie : utilise python-docx pour insérer des paragraphes AVANT
les ancres existantes (article 6.2 pour 6.1 précisé ; article 14
pour les articles 13 bis/ter/quater).
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from copy import deepcopy

ONEDRIVE = Path(r"C:\Users\PC\OneDrive\Documents")
INPUT = ONEDRIVE / "contrat-diara-final1.docx"
OUTPUT = ONEDRIVE / "contrat-diara-final2.docx"

DM_DARK = RGBColor(0x06, 0x4E, 0x3B)
BLACK = RGBColor(0x00, 0x00, 0x00)


def style_run(run, bold=False, italic=False, size=11, color=BLACK, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_paragraph_before(reference_paragraph, text, *,
                         heading=False,
                         bold=False,
                         italic=False,
                         bullet=False,
                         size=11,
                         color=BLACK):
    """
    Insère un nouveau paragraphe AVANT le paragraphe de référence,
    en utilisant le même style de base que le doc.
    """
    new_p = reference_paragraph.insert_paragraph_before("")
    if heading:
        # Heading style : DM_DARK, gras, 13 pt
        run = new_p.add_run(text)
        style_run(run, bold=True, size=13, color=DM_DARK)
        new_p.paragraph_format.space_before = Pt(18)
        new_p.paragraph_format.space_after = Pt(6)
        new_p.paragraph_format.keep_with_next = True
    elif bullet:
        # Bullet list
        new_p.style = "List Bullet"
        run = new_p.add_run(text)
        style_run(run, size=size, color=color)
        new_p.paragraph_format.space_after = Pt(3)
        new_p.paragraph_format.line_spacing = 1.2
    else:
        run = new_p.add_run(text)
        style_run(run, bold=bold, italic=italic, size=size, color=color)
        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        new_p.paragraph_format.space_after = Pt(6)
        new_p.paragraph_format.line_spacing = 1.25
    return new_p


def find_paragraph_by_prefix(doc, prefix):
    """Trouve le premier paragraphe dont le texte commence par prefix."""
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def find_paragraph_by_contains(doc, needle):
    """Trouve le premier paragraphe contenant needle."""
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


# ============================================================================
# CONTENU À INSÉRER
# ============================================================================

PARAGRAPHE_6_1_BIS = (
    "Le Mandant ne saurait être tenu de rétrocéder au Mandataire des "
    "commissions qu'il n'aurait pas effectivement encaissées auprès de ses "
    "clients vendeurs. En cas de défaillance, d'impayé ou de litige avec un "
    "client vendeur empêchant l'encaissement effectif de la commission "
    "d'agence, le Mandant en informera le Mandataire dans un délai de "
    "quinze (15) jours ouvrés à compter de la connaissance de l'événement "
    "et conservera la liberté d'engager, à sa diligence, les actions de "
    "recouvrement appropriées. La rétrocession ne deviendra exigible qu'à "
    "la date de l'encaissement effectif des sommes recouvrées."
)

# Article 13 bis (résiliation par convenance)
ART_13_BIS = [
    ("heading", "ARTICLE 13 BIS — RÉSILIATION ANTICIPÉE PAR CONVENANCE DU MANDANT"),
    ("bold",
     "13 bis.1 — Faculté de résiliation par convenance"),
    ("body",
     "Le Mandant se réserve la faculté de résilier le présent contrat par "
     "anticipation, par convenance et sans motif imputable au Mandataire, à "
     "tout moment, par notification écrite avec accusé de réception au "
     "Mandataire."),
    ("bold",
     "13 bis.2 — Préavis"),
    ("body",
     "Cette résiliation est subordonnée au respect d'un préavis minimal de :"),
    ("bullet",
     "(i) UN (1) mois calendaire si la résiliation intervient avant la première vente effectivement conclue par le Mandataire sous mandat Eurealimmo ;"),
    ("bullet",
     "(ii) TROIS (3) mois calendaires entre la première vente et le douzième mois du contrat ;"),
    ("bullet",
     "(iii) SIX (6) mois calendaires au-delà du douzième mois."),
    ("bold",
     "13 bis.3 — Indemnité forfaitaire avant la première vente"),
    ("body",
     "En cas de résiliation par convenance par le Mandant intervenant AVANT "
     "la première transaction effectivement conclue par le Mandataire sous "
     "mandat Eurealimmo, le Mandant versera au Mandataire, en réparation "
     "intégrale et forfaitaire de l'ensemble des préjudices subis, une "
     "indemnité forfaitaire et transactionnelle de TROIS CENTS EUROS "
     "(300 EUR)."),
    ("body",
     "Cette indemnité est destinée à couvrir intégralement les frais de "
     "transfert administratif du Mandataire vers tout autre Mandant titulaire "
     "de la carte T (formalités CCI, mise à jour RSAC, modification "
     "d'attestation RCP), le Mandataire reconnaissant qu'il dispose déjà, à "
     "la date de signature du présent contrat, d'une formation continue ALUR "
     "à jour, d'une attestation RCP en cours de validité et d'un cadre "
     "comptable opérationnel."),
    ("bold",
     "13 bis.4 — Indemnité après la première vente"),
    ("body",
     "En cas de résiliation par convenance par le Mandant intervenant APRÈS "
     "la première vente conclue sous mandat Eurealimmo, le régime de "
     "l'article 13 (article L134-12 du Code de commerce) s'applique."),
    ("bold",
     "13 bis.5 — Caractère transactionnel"),
    ("body",
     "Le Mandataire reconnaît que l'indemnité forfaitaire prévue à l'article "
     "13 bis.3 couvre intégralement son préjudice dans le cas visé et "
     "constitue une transaction au sens de l'article 2044 du Code civil, à "
     "condition que le préavis prévu à l'article 13 bis.2 (i) soit "
     "effectivement respecté par le Mandant. Le Mandataire renonce, dans ce "
     "cas et à concurrence du préjudice ainsi indemnisé, à toute action "
     "complémentaire fondée sur l'article L134-12 ou l'article L442-1 du "
     "Code de commerce."),
    ("bold",
     "13 bis.6 — Exclusion"),
    ("body",
     "Le présent article 13 bis ne s'applique pas en cas de cessation "
     "résultant : (i) d'une faute du Mandataire (article 13 alinéa 2) ; "
     "(ii) d'une procédure collective ouverte à l'encontre du Mandant ; "
     "(iii) du retrait de la carte T par la CCI pour motif imputable au "
     "Mandant."),
]

# Article 13 ter (mise en sommeil)
ART_13_TER = [
    ("heading", "ARTICLE 13 TER — MISE EN SOMMEIL TEMPORAIRE DU MANDANT"),
    ("bold", "13 ter.1 — Faculté de mise en sommeil"),
    ("body",
     "Le Mandant se réserve la faculté de procéder à une mise en sommeil "
     "de l'activité d'EUREALIMMO SARL pour une durée maximale de DOUZE (12) "
     "mois calendaires, par notification écrite au Mandataire indiquant le "
     "motif (difficultés conjoncturelles, restructuration, événement "
     "personnel du gérant, opération stratégique)."),
    ("bold", "13 ter.2 — Effets de la mise en sommeil"),
    ("body", "Pendant la durée de la mise en sommeil :"),
    ("bullet", "(i) le présent contrat est suspendu et non résilié ;"),
    ("bullet", "(ii) le forfait mensuel de l'article 6.2 cesse d'être dû par le Mandataire ;"),
    ("bullet", "(iii) le verrouillage tarifaire de l'article 6.4 est suspendu, puis reprend à compter de la reprise d'activité notifiée par le Mandant ;"),
    ("bullet", "(iv) la franchise des six (6) mois (article 6.3) non encore consommée est reportée à concurrence du solde restant ;"),
    ("bullet", "(v) les droits prévus aux articles 8 (prime de cession) et 8 bis (droit de préemption) sont suspendus et reprennent à compter de la reprise d'activité."),
    ("bold", "13 ter.3 — Faculté de sortie du Mandataire"),
    ("body",
     "Le Mandataire peut, à tout moment pendant la mise en sommeil, "
     "résilier le contrat sans préavis ni indemnité, par simple "
     "notification écrite, sans préjudice de l'application des clauses 9 "
     "(protection), 10 (confidentialité) et 11 (propriété intellectuelle), "
     "qui demeurent en vigueur pendant les durées qui leur sont propres."),
    ("bold", "13 ter.4 — Reprise d'activité"),
    ("body",
     "Le Mandant notifie au Mandataire la reprise d'activité par écrit. Si "
     "la mise en sommeil dépasse douze (12) mois sans reprise notifiée, le "
     "contrat est résilié de plein droit, sans indemnité de cessation due "
     "à l'une ou l'autre des Parties, le Mandataire conservant néanmoins le "
     "bénéfice de l'indemnité forfaitaire de l'article 13 bis.3 si la "
     "résiliation intervient avant la première vente."),
    ("bold", "13 ter.5 — Limitation"),
    ("body",
     "La présente faculté de mise en sommeil ne peut être invoquée qu'UNE "
     "(1) seule fois au cours de la durée du contrat de trente-six (36) "
     "mois prévue à l'article 2.1."),
]

# Article 13 quater (inactivité prolongée)
ART_13_QUATER = [
    ("heading", "ARTICLE 13 QUATER — INACTIVITÉ PROLONGÉE DU MANDATAIRE"),
    ("body",
     "À défaut, pour le Mandataire, d'avoir conclu sous mandat Eurealimmo "
     "au moins UNE (1) transaction effective OU d'avoir présenté au Mandant "
     "au moins UN (1) référé qualifié au sens de l'article 7.4 ayant abouti "
     "à la signature d'un contrat de mandat avec ce référé, durant les "
     "DOUZE (12) mois calendaires consécutifs suivant la date d'effet du "
     "présent contrat, les Parties pourront convenir de la cessation amiable "
     "du contrat, sans indemnité due au titre de l'article L134-12 du Code "
     "de commerce, après mise en demeure préalable adressée par lettre "
     "recommandée avec accusé de réception, restée infructueuse pendant un "
     "délai de TRENTE (30) jours."),
    ("body",
     "Les Parties reconnaissent que cette inactivité prolongée constitue, "
     "à défaut pour le Mandataire de l'avoir surmontée malgré la mise en "
     "demeure, un motif imputable au Mandataire au sens de l'article 13 "
     "alinéa 2 du présent contrat et de l'article L134-13 du Code de "
     "commerce."),
    ("body",
     "Cette stipulation s'applique sans préjudice des cas de force majeure "
     "légalement reconnus (article 1218 du Code civil), notamment en cas "
     "d'empêchement médical grave, de maternité, paternité ou adoption "
     "documenté, qui suspendent le délai de douze (12) mois pour toute la "
     "durée du fait justificatif."),
]


def insert_block(reference_paragraph, block):
    """Insère un bloc (liste de tuples (type, text)) AVANT le paragraphe de référence."""
    for kind, text in block:
        if kind == "heading":
            add_paragraph_before(reference_paragraph, text, heading=True)
        elif kind == "bold":
            add_paragraph_before(reference_paragraph, text, bold=True)
        elif kind == "body":
            add_paragraph_before(reference_paragraph, text)
        elif kind == "bullet":
            add_paragraph_before(reference_paragraph, text, bullet=True)


def main():
    if not INPUT.exists():
        raise SystemExit(f"Introuvable : {INPUT}")

    doc = Document(str(INPUT))

    # 1) Précision art. 6.1 : insérer avant le titre 6.2
    p_62 = find_paragraph_by_prefix(doc, "6.2")
    if p_62 is None:
        raise SystemExit("Ancre 6.2 non trouvée")
    add_paragraph_before(p_62, PARAGRAPHE_6_1_BIS)
    print("[OK] Insertion alinéa 6.1 bis avant 6.2")

    # 2/3/4) Articles 13 bis + 13 ter + 13 quater : avant ARTICLE 14
    p_14 = find_paragraph_by_prefix(doc, "ARTICLE 14")
    if p_14 is None:
        raise SystemExit("Ancre ARTICLE 14 non trouvée")
    insert_block(p_14, ART_13_BIS)
    print("[OK] Insertion ARTICLE 13 BIS")
    insert_block(p_14, ART_13_TER)
    print("[OK] Insertion ARTICLE 13 TER")
    insert_block(p_14, ART_13_QUATER)
    print("[OK] Insertion ARTICLE 13 QUATER")

    doc.save(str(OUTPUT))
    print(f"\nOK -> {OUTPUT}")


if __name__ == "__main__":
    main()
