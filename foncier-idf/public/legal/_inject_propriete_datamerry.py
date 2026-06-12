"""
Précise dans le contrat final2.docx que le module de notarisation
blockchain est la propriété de DATAMERRY SAS (et non d'Eurealimmo).

Modifie 2 articles :
  - 8.2 bis : élargit la description de la PI DATAMERRY (ajoute le module)
  - 16.6 : remplace "propriété du Mandant" par "propriété de DATAMERRY SAS"
            + clarifie la licence d'usage à Eurealimmo
+ Ajoute 16.6 bis : régime de propriété intellectuelle distinct (module
  DATAMERRY vs hashes mandats Eurealimmo).

Input  : contrat-diara-final2.docx
Output : contrat-diara-final3.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ONEDRIVE = Path(r"C:\Users\PC\OneDrive\Documents")
INPUT = ONEDRIVE / "contrat-diara-final2.docx"
OUTPUT = ONEDRIVE / "contrat-diara-final3.docx"

DM_DARK = RGBColor(0x06, 0x4E, 0x3B)
BLACK = RGBColor(0x00, 0x00, 0x00)


# ---------------------------------------------------------------------------
# Nouveaux contenus
# ---------------------------------------------------------------------------

NEW_8_2_BIS = (
    "Pour lever toute ambiguïté, la cession isolée des parts ou du fonds de "
    "commerce de DATAMERRY SAS, entité juridique distincte d'EUREALIMMO SARL "
    "et titulaire de la plateforme technologique comprenant notamment : "
    "(i) le moteur d'estimation immobilière, (ii) les bases de données "
    "agrégées (DVF, OSM, Wikidata, IDFM PRIM, INSEE, cadastre, permis de "
    "construire), (iii) le module de notarisation cryptographique des "
    "mandats sur blockchain Solana (algorithmes, smart contracts, "
    "infrastructure on-chain, interfaces de génération des certificats), et "
    "(iv) la marque DATAMERRY® et l'ensemble de ses dérivés, ne constitue PAS "
    "un fait générateur de la prime de cession prévue au présent article."
)

NEW_16_6 = (
    "Le module de notarisation blockchain est et restera entièrement GRATUIT "
    "pour le Mandataire pendant toute la durée du présent contrat, y compris "
    "en cas de renouvellement. Le module, les smart contracts associés, leur "
    "code source, leur infrastructure on-chain et leurs développements "
    "ultérieurs demeurent la propriété exclusive et inaliénable de "
    "DATAMERRY SAS, distincte juridiquement d'EUREALIMMO SARL. EUREALIMMO "
    "SARL en bénéficie au titre d'une licence d'usage non exclusive, non "
    "transférable et révocable consentie par DATAMERRY SAS pour les besoins "
    "exclusifs de l'exploitation du réseau Eurealimmo Réseau. DATAMERRY SAS "
    "se réserve la faculté de valoriser, céder ou exploiter commercialement "
    "ledit module auprès de tiers (notamment d'autres cabinets immobiliers, "
    "via le Solana Mobile dApp Store ou par cession à un acquéreur tiers), "
    "sans que cela puisse être assimilé à une cession au sens des articles 8 "
    "et 8.2 bis du présent contrat."
)

# Nouvelle clause 16.6 bis (distinction module / hashes)
NEW_16_6_BIS_TITLE = "16.6 bis — Distinction entre module technique et données ancrées"
NEW_16_6_BIS_BODY = (
    "Les Parties conviennent expressément de la distinction juridique "
    "suivante : (i) le module technique de notarisation (algorithmes, smart "
    "contracts, infrastructure on-chain, interfaces) constitue une propriété "
    "exclusive de DATAMERRY SAS ; (ii) les hashes SHA-256 effectivement "
    "publiés sur la blockchain Solana, dérivés des mandats Hoguet détenus "
    "par EUREALIMMO SARL en sa qualité de titulaire de la carte T, demeurent "
    "attachés à EUREALIMMO SARL sans pour autant affecter la propriété de "
    "DATAMERRY SAS sur le module technique sous-jacent ; (iii) le certificat "
    "de notarisation délivré, le cas échéant, au Mandataire (article 16.3) "
    "est signé et engageant pour EUREALIMMO SARL, seule habilitée au regard "
    "de la loi Hoguet à émettre des actes relatifs aux mandats immobiliers "
    "qu'elle détient."
)


# ---------------------------------------------------------------------------
# Utilitaires
# ---------------------------------------------------------------------------

def replace_paragraph_text(p, new_text):
    """
    Remplace le texte du paragraphe en préservant le style du premier run.
    """
    if not p.runs:
        # Pas de run : ajouter un run vide pour préserver la police
        p.add_run("")
    # Conserver le format du premier run, vider les autres
    first_run = p.runs[0]
    for run in p.runs[1:]:
        run.text = ""
    first_run.text = new_text


def style_run(run, bold=False, italic=False, size=11, color=BLACK, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_paragraph_before(reference_paragraph, text, *, bold=False, body=False):
    new_p = reference_paragraph.insert_paragraph_before("")
    run = new_p.add_run(text)
    if bold:
        style_run(run, bold=True, size=11)
    else:
        style_run(run, size=11)
    if body:
        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        new_p.paragraph_format.space_after = Pt(6)
        new_p.paragraph_format.line_spacing = 1.25
    return new_p


def find_paragraph_by_contains(doc, needle):
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    if not INPUT.exists():
        raise SystemExit(f"Introuvable : {INPUT}")

    doc = Document(str(INPUT))

    # 1) Remplacer le corps de l'art. 8.2 bis (paragraphe contenant
    #    "Pour lever toute ambiguïté")
    p_82bis = find_paragraph_by_contains(doc, "Pour lever toute ambiguïté")
    if p_82bis is None:
        # Essayer sans accent (selon encodage)
        p_82bis = find_paragraph_by_contains(doc, "Pour lever toute ambigu")
    if p_82bis is None:
        raise SystemExit("Ancre 8.2 bis non trouvée")
    replace_paragraph_text(p_82bis, NEW_8_2_BIS)
    print("[OK] Mise à jour 8.2 bis (ajout module notarisation à la PI DATAMERRY)")

    # 2) Remplacer le corps de l'art. 16.6
    p_166 = find_paragraph_by_contains(doc, "Le module de notarisation blockchain est et restera")
    if p_166 is None:
        raise SystemExit("Ancre 16.6 non trouvée")
    replace_paragraph_text(p_166, NEW_16_6)
    print("[OK] Mise à jour 16.6 (propriété DATAMERRY SAS + licence Eurealimmo)")

    # 3) Insérer 16.6 bis juste avant 16.7
    p_167 = find_paragraph_by_contains(doc, "16.7 — Faculté d'expérimentation")
    if p_167 is None:
        p_167 = find_paragraph_by_contains(doc, "16.7")
    if p_167 is None:
        raise SystemExit("Ancre 16.7 non trouvée")
    add_paragraph_before(p_167, NEW_16_6_BIS_TITLE, bold=True)
    add_paragraph_before(p_167, NEW_16_6_BIS_BODY, body=True)
    print("[OK] Insertion 16.6 bis (distinction module / hashes / certificat)")

    doc.save(str(OUTPUT))
    print(f"\nOK -> {OUTPUT}")


if __name__ == "__main__":
    main()
