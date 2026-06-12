"""
Patch contrat V3 → V4 : passe le 15 % Standard de "pendant 12 mois" à "à vie".

Modification ciblée à l'article 7.2 (bareme Associée Fondatrice).

Input  : contrat-mandat-diara-eurealimmo V3.docx
Output : contrat-mandat-diara-eurealimmo V4.docx (+ .pdf)
"""

from pathlib import Path
from docx import Document

ONEDRIVE = Path(r"C:\Users\PC\OneDrive\Documents")
INPUT = ONEDRIVE / "contrat-mandat-diara-eurealimmo V3.docx"
OUTPUT_DOCX = ONEDRIVE / "contrat-mandat-diara-eurealimmo V4.docx"
OUTPUT_PDF = ONEDRIVE / "contrat-mandat-diara-eurealimmo V4.pdf"

# Texte ancien à remplacer (peut avoir des variantes selon les caractères)
OLD_FRAGMENTS = [
    "versées pendant DOUZE (12) mois",
    "versees pendant DOUZE (12) mois",
    "versées pendant douze (12) mois",
]

# Texte nouveau (matche le format HNWI)
NEW_TEXT = "versées À VIE tant que le contrat de mandat du référencé demeure en vigueur"


def patch_paragraph(p):
    """Remplace dans les runs si possible, sinon dans le texte global."""
    full = "".join(r.text for r in p.runs)
    new = full
    for old in OLD_FRAGMENTS:
        if old in new:
            new = new.replace(old, NEW_TEXT)
    if new != full and p.runs:
        first = p.runs[0]
        for r in p.runs[1:]:
            r.text = ""
        first.text = new
        return True
    return False


def main():
    if not INPUT.exists():
        raise SystemExit(f"Source introuvable : {INPUT}")

    doc = Document(str(INPUT))
    patched = 0
    for p in doc.paragraphs:
        if patch_paragraph(p):
            patched += 1
            print(f"OK -> patch dans paragraphe : {p.text[:120]!r}")

    # Aussi dans les tables (rare ici, par sécurité)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if patch_paragraph(p):
                        patched += 1

    if patched == 0:
        print("AVERTISSEMENT : aucun patch effectué. Vérifier le V3.")
    else:
        print(f"Total patchs : {patched}")

    doc.save(str(OUTPUT_DOCX))
    print(f"OK -> {OUTPUT_DOCX}")

    try:
        from docx2pdf import convert
        convert(str(OUTPUT_DOCX), str(OUTPUT_PDF))
        print(f"OK -> {OUTPUT_PDF}")
    except Exception as e:
        print(f"PDF conversion ignorée : {e}")


if __name__ == "__main__":
    main()
