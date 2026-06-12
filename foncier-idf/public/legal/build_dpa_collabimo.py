"""
Génère le DPA (Data Processing Agreement) RGPD art. 28 entre :
  - COLLABIMO (Diara CAMARA) — Responsable de traitement
  - EUREALIMMO SARL + DATAMERRY SAS — Sous-traitants

Format : PDF 4 pages, prêt à signer.

Usage : python build_dpa_collabimo.py
Sortie : dpa-collabimo-datamerry.pdf
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DOCX = Path(r"C:\Users\PC\workspace\idf-foncier\foncier-idf\public\legal\dpa-collabimo-datamerry.docx")
OUTPUT_PDF = Path(r"C:\Users\PC\workspace\idf-foncier\foncier-idf\public\legal\dpa-collabimo-datamerry.pdf")

DM_DARK = RGBColor(0x06, 0x4E, 0x3B)
DM_GREEN = RGBColor(0x10, 0xB9, 0x81)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x66, 0x66, 0x66)


def set_run(run, bold=False, italic=False, size=11, color=BLACK, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_heading(doc, text, level=1, color=DM_DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    size = 16 if level == 0 else (13 if level == 1 else 11)
    set_run(run, bold=True, size=size, color=color)
    return p


def add_body(doc, text, bold=False, italic=False, align="justify", size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run(run, bold=bold, italic=italic, size=size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_run(run)
    return p


def main():
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Style par défaut
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ============= PAGE DE GARDE =============
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("CONVENTION DE SOUS-TRAITANCE")
    set_run(r, bold=True, size=18, color=DM_DARK)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("DES DONNÉES PERSONNELLES")
    set_run(r, bold=True, size=14, color=DM_DARK)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.paragraph_format.space_after = Pt(8)
    r = sub2.add_run("(Data Processing Agreement — DPA)")
    set_run(r, italic=True, size=12, color=GREY)

    sub3 = doc.add_paragraph()
    sub3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub3.paragraph_format.space_after = Pt(18)
    r = sub3.add_run(
        "Conforme à l'article 28 du Règlement (UE) 2016/679 (RGPD)"
    )
    set_run(r, italic=True, size=10, color=GREY)

    add_body(doc, "Entre les soussignés :", bold=True)

    # ---- Responsable de traitement ----
    add_heading(doc, "Le Responsable de traitement — COLLABIMO", level=2)
    add_body(doc,
        "Mme Diara CAMARA, exerçant sous l'enseigne commerciale COLLABIMO, "
        "agissant en sa qualité de responsable de traitement au sens de "
        "l'article 4-7) du RGPD ; ci-après dénommée « le Responsable de "
        "traitement ».",
        italic=True)

    # ---- Sous-traitants ----
    add_heading(doc, "Les Sous-traitants", level=2)
    add_body(doc,
        "EUREALIMMO SARL, société à responsabilité limitée, capital social "
        "100 €, siège social Paris (75008), SIREN 984 449 470, titulaire de "
        "la carte professionnelle T n° CPI 7501 2024 000 219, représentée par "
        "M. Samuel BRUNO en sa qualité de gérant ;",
        italic=True)
    add_body(doc,
        "Et DATAMERRY SAS (société en cours de constitution à la date des "
        "présentes — à défaut M. Samuel BRUNO à titre personnel), titulaire "
        "de la plateforme technologique DATAMERRY®, agissant en qualité de "
        "sous-traitants au sens de l'article 4-8) du RGPD ; ci-après "
        "ensemble dénommés « les Sous-traitants ».",
        italic=True)

    add_body(doc, "Ensemble dénommés « les Parties ».", italic=True)
    doc.add_page_break()

    # ============= ARTICLE 1 — OBJET =============
    add_heading(doc, "ARTICLE 1 — OBJET ET CHAMP D'APPLICATION", level=1)
    add_body(doc,
        "La présente Convention de sous-traitance des données (ci-après le "
        "« DPA ») a pour objet de définir les conditions dans lesquelles "
        "les Sous-traitants s'engagent à effectuer, pour le compte du "
        "Responsable de traitement, les opérations de traitement de "
        "données à caractère personnel définies ci-après.")
    add_body(doc,
        "Elle s'applique exclusivement aux données traitées par les "
        "Sous-traitants dans le cadre de la fourniture de la plateforme "
        "technologique DATAMERRY® au Responsable de traitement (estimation "
        "immobilière, capture de leads, génération de rapports, "
        "notarisation cryptographique).")

    # ============= ARTICLE 2 — FINALITÉS ET CATÉGORIES =============
    add_heading(doc, "ARTICLE 2 — FINALITÉS ET CATÉGORIES DE TRAITEMENT", level=1)
    add_body(doc, "2.1 — Finalités du traitement", bold=True)
    add_body(doc, "Les Sous-traitants traitent les données aux seules fins suivantes :")
    add_bullet(doc, "Hébergement et stockage des leads capturés via la plateforme ;")
    add_bullet(doc, "Calcul d'estimations immobilières via les modèles statistiques ;")
    add_bullet(doc, "Génération de rapports PDF brandés Collabimo ;")
    add_bullet(doc, "Restitution des données au Responsable de traitement via le tableau de bord ;")
    add_bullet(doc, "Notarisation cryptographique des opérations à fins de traçabilité.")

    add_body(doc, "2.2 — Catégories de données traitées", bold=True)
    add_bullet(doc, "Identité : nom, prénom du prospect vendeur ;")
    add_bullet(doc, "Coordonnées : email, téléphone, adresse postale ;")
    add_bullet(doc, "Données techniques : adresse du bien à estimer, surface, type, pièces ;")
    add_bullet(doc, "Données de connexion : adresse IP, user-agent, horodatage de capture ;")
    add_bullet(doc, "Données d'estimation : prix calculé, comparables, indicateurs.")

    add_body(doc, "2.3 — Catégories de personnes concernées", bold=True)
    add_body(doc,
        "Particuliers ou personnes morales utilisatrices du formulaire "
        "d'estimation immobilière hébergé sur les pages brandées Collabimo "
        "de la plateforme DATAMERRY (typiquement, des vendeurs immobiliers "
        "potentiels).")

    add_body(doc, "2.4 — Durée du traitement", bold=True)
    add_body(doc,
        "La durée du traitement est strictement limitée à la durée du "
        "contrat de prestation de services entre le Responsable de "
        "traitement et les Sous-traitants, prolongée de TRENTE (30) jours "
        "calendaires au maximum pour les opérations de restitution et de "
        "suppression définitive.")

    doc.add_page_break()

    # ============= ARTICLE 3 — OBLIGATIONS DES SOUS-TRAITANTS =============
    add_heading(doc, "ARTICLE 3 — OBLIGATIONS DES SOUS-TRAITANTS", level=1)
    add_body(doc, "Les Sous-traitants s'engagent à :")
    add_bullet(doc, "Ne traiter les données que sur instruction documentée du Responsable de traitement, conformément à l'article 28-3-a) du RGPD ;")
    add_bullet(doc, "Garantir que les personnes autorisées à traiter les données s'engagent à respecter la confidentialité (article 28-3-b) ;")
    add_bullet(doc, "Mettre en œuvre des mesures techniques et organisationnelles appropriées pour assurer la sécurité du traitement (article 28-3-c et 32) — voir Annexe 1 ;")
    add_bullet(doc, "Notifier au Responsable de traitement, sans délai et au plus tard SOIXANTE-DOUZE (72) heures après en avoir pris connaissance, toute violation de données à caractère personnel (article 33) ;")
    add_bullet(doc, "Aider le Responsable de traitement à répondre aux demandes d'exercice des droits des personnes concernées (accès, rectification, effacement, portabilité, opposition — articles 12 à 22) ;")
    add_bullet(doc, "Mettre à disposition du Responsable de traitement toutes les informations nécessaires à la démonstration du respect du présent DPA et permettre la réalisation d'audits, y compris des inspections.")

    # ============= ARTICLE 4 — SÉCURITÉ =============
    add_heading(doc, "ARTICLE 4 — MESURES DE SÉCURITÉ", level=1)
    add_body(doc,
        "Les Sous-traitants mettent en œuvre les mesures techniques et "
        "organisationnelles suivantes (article 32 RGPD) :")
    add_bullet(doc, "Chiffrement des données en transit (TLS 1.3 minimum) et au repos (AES-256) ;")
    add_bullet(doc, "Séparation physique entre la base contenant les données personnelles (PII) et la base des données anonymisées utilisée pour les modèles statistiques ;")
    add_bullet(doc, "Contrôle d'accès strict par authentification forte (magic link Resend + cookies HttpOnly sécurisés) et journalisation de tous les accès aux PII ;")
    add_bullet(doc, "Row Level Security (RLS) activé au niveau de la base de données pour empêcher l'accès croisé entre cabinets ;")
    add_bullet(doc, "Sauvegardes quotidiennes chiffrées avec rétention 30 jours ;")
    add_bullet(doc, "Mises à jour de sécurité régulières des dépendances logicielles ;")
    add_bullet(doc, "K-anonymisation (k≥5) appliquée à toute statistique agrégée publiée.")

    # ============= ARTICLE 5 — LOCALISATION =============
    add_heading(doc, "ARTICLE 5 — LOCALISATION ET TRANSFERTS", level=1)
    add_body(doc,
        "Les données sont hébergées exclusivement au sein de l'Union "
        "européenne :")
    add_bullet(doc, "Base de données Supabase (PostgreSQL) — Region Frankfurt (DE), conforme RGPD ;")
    add_bullet(doc, "Hébergement applicatif Vercel — Region Frankfurt (DE), conforme RGPD ;")
    add_bullet(doc, "Service d'envoi d'emails Resend — Region Dublin (IE), conforme RGPD.")
    add_body(doc,
        "Aucun transfert de données vers un pays tiers (hors UE/EEE) ne "
        "sera effectué sans accord écrit préalable du Responsable de "
        "traitement et sans mise en place des garanties appropriées "
        "(article 44 et suivants RGPD).")

    # ============= ARTICLE 6 — SOUS-TRAITANCE ULTÉRIEURE =============
    add_heading(doc, "ARTICLE 6 — SOUS-TRAITANCE ULTÉRIEURE", level=1)
    add_body(doc,
        "Les Sous-traitants peuvent recourir aux sous-traitants ultérieurs "
        "suivants, autorisés à la date de signature :")
    add_bullet(doc, "Supabase Inc. — hébergement base de données PostgreSQL (Frankfurt, UE) ;")
    add_bullet(doc, "Vercel Inc. — hébergement applicatif Next.js (Frankfurt, UE) ;")
    add_bullet(doc, "Resend Inc. — service d'envoi transactionnel (Dublin, UE).")
    add_body(doc,
        "Toute modification de la liste des sous-traitants ultérieurs sera "
        "préalablement notifiée au Responsable de traitement par écrit, "
        "qui disposera d'un délai de QUINZE (15) jours pour s'y opposer.")

    doc.add_page_break()

    # ============= ARTICLE 7 — DROITS DU RESPONSABLE =============
    add_heading(doc, "ARTICLE 7 — DROITS DU RESPONSABLE DE TRAITEMENT", level=1)
    add_body(doc, "Le Responsable de traitement peut, à tout moment :")
    add_bullet(doc, "Demander un export complet de l'intégralité des données au format CSV ou JSON, fourni sous QUARANTE-HUIT (48) heures ouvrées ;")
    add_bullet(doc, "Demander la rectification ou la suppression de tout ou partie des données ;")
    add_bullet(doc, "Réaliser un audit annuel des mesures de sécurité, sur rendez-vous et sous accord de confidentialité ;")
    add_bullet(doc, "Obtenir un rapport mensuel de journalisation des accès aux données personnelles ;")
    add_bullet(doc, "Vérifier publiquement l'ancrage cryptographique mensuel des opérations sur le registre blockchain choisi par les Sous-traitants.")

    # ============= ARTICLE 8 — SORT EN FIN DE CONTRAT =============
    add_heading(doc, "ARTICLE 8 — SORT DES DONNÉES EN FIN DE CONTRAT", level=1)
    add_body(doc,
        "À la cessation, pour quelque cause que ce soit, du contrat de "
        "prestation de services :")
    add_bullet(doc, "Les Sous-traitants exportent, sous TRENTE (30) jours calendaires, l'intégralité des données traitées pour le compte du Responsable de traitement au format CSV ou JSON structuré ;")
    add_bullet(doc, "Les Sous-traitants suppriment définitivement, sous le même délai, l'ensemble des données personnelles de leurs serveurs et des sauvegardes, à l'exception des données anonymisées (au sens du paragraphe suivant) ;")
    add_bullet(doc, "Les données anonymisées au sens des critères CNIL WP216 (non-individualisation, non-corrélation, non-inférence) peuvent être conservées par les Sous-traitants à des fins statistiques, sans limitation de durée et sans préjudice pour le Responsable de traitement ou les personnes concernées.")

    # ============= ARTICLE 9 — VIOLATION DE DONNÉES =============
    add_heading(doc, "ARTICLE 9 — VIOLATION DE DONNÉES", level=1)
    add_body(doc,
        "En cas de violation de données à caractère personnel, les "
        "Sous-traitants s'engagent à :")
    add_bullet(doc, "Notifier le Responsable de traitement sans délai et au plus tard 72h après en avoir pris connaissance, en précisant la nature de la violation, les catégories et nombre approximatif de personnes concernées, les conséquences probables et les mesures prises ;")
    add_bullet(doc, "Documenter la violation et fournir au Responsable de traitement tous éléments lui permettant, le cas échéant, de notifier l'incident à la CNIL et aux personnes concernées (articles 33 et 34 RGPD).")

    # ============= ARTICLE 10 — RESPONSABILITÉ =============
    add_heading(doc, "ARTICLE 10 — RESPONSABILITÉ", level=1)
    add_body(doc,
        "Chaque Partie est responsable des dommages causés par tout "
        "manquement de sa part au présent DPA ou au RGPD. La responsabilité "
        "des Sous-traitants est plafonnée, sauf faute intentionnelle ou "
        "lourde, à un montant équivalent à DIX (10) fois le coût "
        "d'abonnement annuel net facturé au Responsable de traitement au "
        "titre du service de plateforme DATAMERRY®.")

    # ============= ARTICLE 11 — DURÉE ET RÉSILIATION =============
    add_heading(doc, "ARTICLE 11 — DURÉE", level=1)
    add_body(doc,
        "Le présent DPA prend effet à la date de signature et demeure en "
        "vigueur pendant toute la durée du contrat de prestation de "
        "services entre le Responsable de traitement et les Sous-traitants, "
        "puis pendant TRENTE (30) jours additionnels pour les opérations "
        "de restitution et de suppression définitive prévues à l'article 8.")

    # ============= ARTICLE 12 — DROIT APPLICABLE =============
    add_heading(doc, "ARTICLE 12 — DROIT APPLICABLE ET JURIDICTION", level=1)
    add_body(doc,
        "Le présent DPA est régi par le droit français et par le "
        "Règlement (UE) 2016/679 (RGPD). Tout différend relatif à son "
        "interprétation ou son exécution sera, à défaut d'accord amiable, "
        "soumis à la compétence exclusive du Tribunal de commerce de Paris.")

    # ============= SIGNATURES =============
    doc.add_page_break()
    add_heading(doc, "SIGNATURES", level=1)
    add_body(doc,
        "Fait à _________________, le _____ / _____ / 2026, en deux (2) "
        "exemplaires originaux.")
    add_body(doc, "")
    add_body(doc, "")

    # 2 colonnes signatures
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(40)
    run = p.add_run("Le Responsable de traitement\n\n")
    set_run(run, bold=True, size=11)
    run = p.add_run("Mme Diara CAMARA\nCollabimo\n\n\n\n_____________________________\n(Signature précédée de « Lu et approuvé »)")
    set_run(run, size=10, color=GREY)

    p = doc.add_paragraph()
    run = p.add_run("Les Sous-traitants\n\n")
    set_run(run, bold=True, size=11)
    run = p.add_run("M. Samuel BRUNO\nEUREALIMMO SARL & DATAMERRY SAS (en formation)\n\n\n\n_____________________________\n(Signature précédée de « Lu et approuvé »)")
    set_run(run, size=10, color=GREY)

    # Save DOCX
    doc.save(str(OUTPUT_DOCX))
    print(f"OK -> {OUTPUT_DOCX}")

    # Conversion PDF
    from docx2pdf import convert
    convert(str(OUTPUT_DOCX), str(OUTPUT_PDF))
    print(f"OK -> {OUTPUT_PDF}")


if __name__ == "__main__":
    main()
