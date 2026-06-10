"""
Génère 4 templates de mandat Hoguet conformes décret 72-678 art. 73,
avec tags docxtemplater {champ} pour remplissage automatique depuis
dim_cabinet_leads.

Templates générés :
  1. mandat-vente.template.docx                — mandat de vente immo
  2. mandat-recherche-acquereur.template.docx  — recherche acquéreur
  3. mandat-mise-en-location.template.docx     — recherche locataire (bailleur)
  4. mandat-recherche-bien-locatif.template.docx — recherche bien à louer

Conformité légale intégrée :
  * Loi Hoguet n° 70-9 art. 6 (mandat écrit obligatoire)
  * Décret 72-678 art. 72 à 78 (forme, durée, prix, dénonciation)
  * Décret 72-678 art. 73 (9 MENTIONS OBLIGATOIRES) :
      1. Identité titulaire carte T + n° carte
      2. N° de registre des mandats
      3. Identité du mandant
      4. Bien concerné
      5. Prix ou budget
      6. Modalités de rémunération (qui paie, TTC)
      7. Durée + conditions de dénonciation
      8. Conditions de renouvellement
      9. Information rétractation L271-1 si signature à distance
  * Code consommation L. 224-25-1 (rétractation 14 j si signature à distance)
  * Loi ALUR 2014 (information précontractuelle)

ATTENTION JURIDIQUE :
  Ces templates sont conformes aux textes en vigueur au 2026-06.
  Il est néanmoins RECOMMANDÉ de les faire valider par un avocat
  immobilier avant la 1re signature en production.
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = Path(__file__).parent / "templates"
OUTPUT_DIR.mkdir(exist_ok=True)

DM_DARK = RGBColor(0x06, 0x4E, 0x3B)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x66, 0x66, 0x66)
GOLD = RGBColor(0xC8, 0xA2, 0x5D)


# ============================================================
# Helpers de mise en forme
# ============================================================

def set_run(run, bold=False, italic=False, size=11, color=BLACK, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_title(doc, text, size=18, color=DM_DARK):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, bold=True, size=size, color=color)
    return p


def add_subtitle(doc, text, size=11, color=GREY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    set_run(r, italic=True, size=size, color=color)
    return p


def add_heading(doc, text, level=1, color=DM_DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    size = 14 if level == 1 else 11
    set_run(r, bold=True, size=size, color=color)
    return p


def add_body(doc, text, bold=False, italic=False, align="justify", size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    set_run(r, bold=bold, italic=italic, size=size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_run(r)
    return p


def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)


def common_header(doc, mandat_type_label):
    """En-tête commun aux 4 templates."""
    add_title(doc, f"MANDAT {mandat_type_label.upper()}")
    add_subtitle(
        doc,
        "Loi Hoguet n° 70-9 du 2 janvier 1970 · Décret 72-678 modifié",
    )

    add_body(
        doc,
        "N° de registre des mandats : {numero_registre}    "
        "Date : {date_mandat}",
        align="center",
        bold=True,
    )

    add_body(doc, "")  # spacer


def common_parties(doc):
    """Identification des parties : Mandant + Mandataire client."""
    add_heading(doc, "LES PARTIES", level=1)

    add_body(doc,
        "Entre les soussignés :",
        bold=True,
    )

    add_body(doc, "LE MANDANT (Titulaire de la carte T)", bold=True)
    add_body(doc,
        "EUREALIMMO SARL, société à responsabilité limitée au capital "
        "de 100 EUR, dont le siège social est situé Paris (75008), "
        "immatriculée au RCS de Paris sous le numéro 984 449 470, "
        "SIRET 98444947000017, représentée par M. Samuel BRUNO en sa "
        "qualité de gérant, titulaire de la carte professionnelle de "
        "transactions sur immeubles et fonds de commerce n° CPI 7501 "
        "2024 000 219 délivrée par la CCI Paris Île-de-France le "
        "[date délivrance], mention « Transactions », sans maniement "
        "de fonds (déclaration CCI Paris au sens de l'article 3-2 de "
        "la loi Hoguet — aucune garantie financière requise) ; "
        "couverte par une assurance de responsabilité civile "
        "professionnelle.",
        align="justify",
    )

    add_body(doc,
        "Représentée pour l'exécution du présent mandat par son agent "
        "commercial : {prenom_mandataire} {nom_mandataire}, immatriculé "
        "au Registre Spécial des Agents Commerciaux du Tribunal de "
        "Commerce de Paris sous le n° {numero_rsac_mandataire}, "
        "habilité par attestation CCI Paris Île-de-France n° "
        "{numero_attestation} (article 9 du décret 72-678).",
        align="justify",
    )

    add_body(doc,
        "Ci-après dénommée « l'Agence » ou « le Mandataire (titulaire) »,",
        italic=True,
    )
    add_body(doc, "D'une part,")

    add_body(doc, "LE CLIENT (Mandant économique au sens du présent mandat)",
             bold=True)
    add_body(doc,
        "{client_civilite} {client_nom} {client_prenom}, "
        "né(e) le {client_date_naissance} à {client_lieu_naissance}, "
        "de nationalité {client_nationalite}, "
        "demeurant {client_adresse}, "
        "joignable au {client_telephone} / {client_email}.",
        align="justify",
    )

    add_body(doc,
        "Ci-après dénommé(e) « le Client » ou « le Mandant économique »,",
        italic=True,
    )
    add_body(doc, "D'autre part,")

    add_body(doc, "Ensemble dénommés « les Parties ».", italic=True)


def common_etat_civil_complement(doc):
    """Bloc à compléter manuellement (état civil du client)."""
    add_heading(doc, "INFORMATIONS COMPLÉMENTAIRES SUR LE CLIENT", level=2)
    add_body(doc, "Régime matrimonial : ____________________________________")
    add_body(doc, "Situation familiale : ____________________________________")
    add_body(doc, "Profession : __________________________________________")


def common_durée_dénonciation(doc):
    """Bloc durée + conditions de dénonciation (art. 73 al. 7)."""
    add_heading(doc, "DURÉE ET DÉNONCIATION DU MANDAT", level=1)

    add_body(doc,
        "Le présent mandat est conclu pour une durée déterminée de "
        "{duree_mois} mois calendaires, à compter de la date de "
        "signature, soit jusqu'au {date_fin}.",
        align="justify",
    )

    add_body(doc,
        "À l'issue de cette durée, le mandat se renouvellera par "
        "tacite reconduction pour des périodes successives d'un (1) "
        "mois, sauf dénonciation par l'une ou l'autre des Parties "
        "par lettre recommandée avec accusé de réception adressée "
        "au plus tard QUINZE (15) jours avant l'échéance.",
        align="justify",
    )

    add_body(doc,
        "Conformément à l'article 78 du décret n° 72-678, le présent "
        "mandat ne pourra en aucun cas être conclu pour une durée "
        "supérieure à TRENTE (30) mois ou être tacitement reconduit "
        "au-delà de ce terme.",
        align="justify",
    )

    add_body(doc,
        "Chaque Partie peut dénoncer le mandat à tout moment, par "
        "lettre recommandée avec accusé de réception adressée à "
        "l'autre Partie, moyennant un préavis de QUINZE (15) jours "
        "calendaires. La dénonciation prend effet à l'expiration du "
        "préavis.",
        align="justify",
    )


def common_renonciation_l271(doc, signature_distance=True):
    """Information sur le droit de rétractation (Code consommation art. L. 271-1 et L. 224-25-1)."""
    add_heading(doc, "INFORMATION SUR LE DROIT DE RÉTRACTATION", level=1)

    add_body(doc,
        "Conformément à l'article L. 224-25-1 du Code de la "
        "consommation, lorsque le présent mandat est conclu HORS "
        "ÉTABLISSEMENT ou à DISTANCE (notamment par signature "
        "électronique sans présence simultanée des Parties), le "
        "Client bénéficie d'un droit de rétractation qu'il peut "
        "exercer dans un délai de QUATORZE (14) jours calendaires "
        "à compter du lendemain du jour de la signature du présent "
        "mandat.",
        align="justify",
    )

    add_body(doc,
        "Ce droit s'exerce par envoi à l'Agence du formulaire de "
        "rétractation joint en annexe, ou par toute autre "
        "déclaration écrite dénuée d'ambiguïté exprimant la volonté "
        "de se rétracter, à l'adresse électronique "
        "contact@eurealimmo.com ou par lettre recommandée avec "
        "accusé de réception au siège social du Mandant.",
        align="justify",
    )

    add_body(doc,
        "Conformément à l'article L. 271-1 du Code de la "
        "construction et de l'habitation, lorsque le présent mandat "
        "porte sur l'acquisition d'un immeuble d'habitation, le "
        "Client non-professionnel dispose en outre d'un délai de "
        "rétractation de DIX (10) jours à compter de la notification "
        "de l'acte de promesse de vente, distinct du droit de "
        "rétractation prévu au présent article.",
        align="justify",
    )


def common_signature_block(doc):
    add_heading(doc, "SIGNATURES", level=1)
    add_body(doc,
        "Fait à _________________________, le {date_mandat}, "
        "en deux (2) exemplaires originaux.",
    )
    add_body(doc, "")
    add_body(doc, "")

    # 2 colonnes signatures
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(50)
    r = p.add_run("Pour l'Agence (Mandataire titulaire carte T)\n\n")
    set_run(r, bold=True, size=11)
    r = p.add_run(
        "{prenom_mandataire} {nom_mandataire}\n"
        "Agent commercial RSAC n° {numero_rsac_mandataire}\n"
        "Pour le compte de EUREALIMMO SARL\n\n\n\n"
        "_____________________________\n"
        "(Signature précédée de « Lu et approuvé »)"
    )
    set_run(r, size=10, color=GREY)

    p = doc.add_paragraph()
    r = p.add_run("Le Client\n\n")
    set_run(r, bold=True, size=11)
    r = p.add_run(
        "{client_civilite} {client_nom} {client_prenom}\n\n\n\n"
        "_____________________________\n"
        "(Signature précédée de « Lu et approuvé »)"
    )
    set_run(r, size=10, color=GREY)


def common_footer(doc):
    add_body(doc, "")
    add_body(doc,
        "EUREALIMMO SARL — SIREN 984 449 470 — Carte T CPI 7501 2024 "
        "000 219 — Sans maniement de fonds. Le présent mandat est "
        "inscrit au registre des mandats du titulaire sous le n° "
        "{numero_registre}, conformément à l'article 72 du décret "
        "72-678 du 20 juillet 1972.",
        italic=True, size=8, align="center",
    )


# ============================================================
# Template 1 : MANDAT DE VENTE
# ============================================================

def build_mandat_vente():
    doc = Document()
    set_margins(doc)

    common_header(doc, "DE VENTE IMMOBILIÈRE")
    common_parties(doc)
    common_etat_civil_complement(doc)

    # Objet du mandat
    add_heading(doc, "ARTICLE 1 — OBJET DU MANDAT", level=1)
    add_body(doc,
        "Le Client donne mandat à l'Agence, qui l'accepte, à l'effet "
        "de rechercher un acquéreur pour le bien immobilier suivant, "
        "dont le Client déclare être pleinement propriétaire et avoir "
        "tous pouvoirs pour le vendre :",
        align="justify",
    )
    add_bullet(doc, "Nature du bien : {bien_type}")
    add_bullet(doc, "Adresse du bien : {bien_adresse}")
    add_bullet(doc, "Surface habitable / utile : {bien_surface} m²")
    add_bullet(doc, "Désignation cadastrale / référence : {bien_reference}")
    add_bullet(doc, "Description sommaire : {bien_description}")

    # Modalité
    add_heading(doc, "ARTICLE 2 — MODALITÉ DU MANDAT", level=1)
    add_body(doc,
        "Le présent mandat est conclu sous la forme : {mandat_modalite}.",
        align="justify",
    )
    add_body(doc,
        "S'il s'agit d'un mandat EXCLUSIF, le Client s'engage à ne "
        "confier la vente du bien décrit ci-dessus à aucun autre "
        "professionnel ni à le vendre par ses propres moyens pendant "
        "toute la durée du mandat. Toute infraction à cette clause "
        "ouvrira droit pour l'Agence à une indemnité forfaitaire "
        "égale à la commission stipulée à l'article 4 ci-après.",
        italic=True, align="justify",
    )

    # Prix net vendeur
    add_heading(doc, "ARTICLE 3 — PRIX DEMANDÉ", level=1)
    add_body(doc,
        "Le Client charge l'Agence de proposer la vente du bien au prix "
        "NET VENDEUR de {prix_net_vendeur} EUR (en lettres : "
        "{prix_net_vendeur_lettres}).",
        align="justify", bold=True,
    )
    add_body(doc,
        "Ce prix s'entend hors honoraires de l'Agence, lesquels sont "
        "à la charge de l'ACQUÉREUR conformément à l'article 4 "
        "ci-après. Le prix affiché à la commercialisation correspondra "
        "donc au prix net vendeur majoré desdits honoraires (« prix "
        "FAI » — frais d'agence inclus).",
        align="justify",
    )

    # Rémunération (art. 73 obligation centrale)
    add_heading(doc, "ARTICLE 4 — HONORAIRES DE L'AGENCE", level=1)
    add_body(doc,
        "Les honoraires de l'Agence sont fixés à "
        "{commission_pct} % TTC du prix net vendeur stipulé à "
        "l'article 3, soit la somme de {commission_eur} EUR TTC "
        "(en lettres : {commission_eur_lettres}).",
        align="justify", bold=True,
    )
    add_body(doc,
        "Ces honoraires sont mis à la CHARGE EXCLUSIVE DE L'ACQUÉREUR, "
        "conformément à l'article 6 alinéa 3 de la loi Hoguet et à "
        "l'arrêté du 10 janvier 2017. Ils ne seront dus que dans le "
        "cas d'une vente effectivement réalisée par l'intermédiaire "
        "de l'Agence et constatée par acte authentique reçu par "
        "notaire.",
        align="justify",
    )

    # Durée + dénonciation
    common_durée_dénonciation(doc)

    # Diagnostics + L. 271-1
    add_heading(doc, "ARTICLE 6 — OBLIGATIONS DU CLIENT", level=1)
    add_bullet(doc,
        "Fournir l'ensemble des diagnostics techniques obligatoires "
        "(DPE, amiante, plomb, gaz, électricité, ERP, etc.) en cours "
        "de validité ;")
    add_bullet(doc,
        "Communiquer à l'Agence l'ensemble des informations matérielles "
        "et juridiques relatives au bien (titre de propriété, copropriété, "
        "procès-verbaux d'AG des 3 dernières années, règlement de "
        "copropriété, charges courantes, travaux votés) ;")
    add_bullet(doc,
        "Informer sans délai l'Agence de toute évolution affectant le "
        "bien ou la transaction (offre directe, retrait, sinistre, etc.) ;")
    add_bullet(doc,
        "Permettre l'organisation des visites dans des conditions "
        "normales et lors de plages horaires raisonnables.")

    # L. 271-1 rétractation
    common_renonciation_l271(doc, signature_distance=True)

    # Données personnelles RGPD
    add_heading(doc, "ARTICLE 8 — PROTECTION DES DONNÉES PERSONNELLES", level=1)
    add_body(doc,
        "Les données personnelles communiquées par le Client sont "
        "collectées et traitées par l'Agence pour les besoins "
        "exclusifs de l'exécution du présent mandat, conformément au "
        "Règlement (UE) 2016/679 (RGPD) et à la loi n° 78-17 du 6 "
        "janvier 1978 modifiée. Le Client peut exercer ses droits "
        "d'accès, de rectification, de suppression et d'opposition "
        "auprès de contact@eurealimmo.com.",
        align="justify",
    )

    # Litiges
    add_heading(doc, "ARTICLE 9 — LITIGES ET MÉDIATION", level=1)
    add_body(doc,
        "En cas de litige relatif à l'exécution ou à l'interprétation "
        "du présent mandat, les Parties s'engagent à rechercher une "
        "solution amiable. À défaut d'accord dans un délai de TRENTE "
        "(30) jours, le Client peut saisir gratuitement le médiateur "
        "de la consommation désigné par l'Agence, dont les coordonnées "
        "sont communiquées sur le site internet du Mandant. À défaut "
        "de règlement amiable, le litige sera porté devant les "
        "juridictions compétentes du ressort du domicile du Client.",
        align="justify",
    )

    common_signature_block(doc)
    common_footer(doc)

    # Annexe formulaire rétractation
    doc.add_page_break()
    add_heading(doc, "ANNEXE — FORMULAIRE DE RÉTRACTATION", level=1)
    add_body(doc,
        "(À compléter et renvoyer uniquement si vous souhaitez vous "
        "rétracter dans les 14 jours suivant la signature)",
        italic=True,
    )
    add_body(doc, "")
    add_body(doc, "À l'attention de : EUREALIMMO SARL, Paris (75008)")
    add_body(doc, "Email : contact@eurealimmo.com")
    add_body(doc, "")
    add_body(doc,
        "Je / Nous (*) ____________________________________ vous "
        "notifie / notifions (*) ma / notre (*) rétractation du "
        "mandat conclu le {date_mandat} et portant sur la vente du "
        "bien sis {bien_adresse}.",
    )
    add_body(doc, "")
    add_body(doc, "Date : ___________________  Signature : ___________________")
    add_body(doc, "")
    add_body(doc, "(*) Rayer la mention inutile.", italic=True, size=9)

    return doc


# ============================================================
# Template 2 : MANDAT DE RECHERCHE ACQUÉREUR
# ============================================================

def build_mandat_recherche_acquereur():
    doc = Document()
    set_margins(doc)

    common_header(doc, "DE RECHERCHE D'UN ACQUÉREUR")
    common_parties(doc)
    common_etat_civil_complement(doc)

    add_heading(doc, "ARTICLE 1 — OBJET DU MANDAT", level=1)
    add_body(doc,
        "Le Client donne mandat à l'Agence, qui l'accepte, à l'effet "
        "de rechercher pour son compte un bien immobilier répondant "
        "aux critères suivants, en vue de son acquisition :",
        align="justify",
    )
    add_bullet(doc, "Type de bien recherché : {bien_type}")
    add_bullet(doc, "Zone géographique : {bien_zone_geo}")
    add_bullet(doc, "Surface minimale : {bien_surface_min} m²")
    add_bullet(doc, "Surface maximale : {bien_surface_max} m²")
    add_bullet(doc, "Nombre de pièces : {bien_nb_pieces}")
    add_bullet(doc, "Critères complémentaires : {bien_criteres_complementaires}")

    add_heading(doc, "ARTICLE 2 — BUDGET MAXIMUM", level=1)
    add_body(doc,
        "Le Client fixe son budget maximum d'acquisition (frais "
        "d'agence et frais notariés inclus) à : {prix_max} EUR "
        "(en lettres : {prix_max_lettres}).",
        align="justify", bold=True,
    )

    add_heading(doc, "ARTICLE 3 — HONORAIRES DE L'AGENCE", level=1)
    add_body(doc,
        "Les honoraires de l'Agence sont fixés à "
        "{commission_pct} % TTC du prix d'acquisition effectif du "
        "bien acquis par l'intermédiaire de l'Agence, soit un "
        "montant maximum estimatif de {commission_eur} EUR TTC.",
        align="justify", bold=True,
    )
    add_body(doc,
        "Ces honoraires sont mis à la CHARGE EXCLUSIVE DU CLIENT "
        "(acquéreur), conformément à l'article 6 alinéa 3 de la loi "
        "Hoguet. Ils ne seront dus que dans le cas d'une acquisition "
        "effectivement réalisée par l'intermédiaire de l'Agence et "
        "constatée par acte authentique reçu par notaire.",
        align="justify",
    )

    common_durée_dénonciation(doc)

    add_heading(doc, "ARTICLE 5 — OBLIGATIONS DU CLIENT", level=1)
    add_bullet(doc, "Communiquer à l'Agence des critères de recherche précis ;")
    add_bullet(doc, "Avertir l'Agence de toute évolution du projet ;")
    add_bullet(doc, "Effectuer les visites dans des conditions courtoises ;")
    add_bullet(doc,
        "S'abstenir de contacter directement un vendeur dont les "
        "coordonnées ont été révélées par l'Agence dans le cadre du "
        "présent mandat.")

    common_renonciation_l271(doc, signature_distance=True)

    add_heading(doc, "ARTICLE 7 — PROTECTION DES DONNÉES PERSONNELLES", level=1)
    add_body(doc,
        "Les données personnelles communiquées par le Client sont "
        "collectées et traitées par l'Agence conformément au RGPD. "
        "Le Client peut exercer ses droits auprès de "
        "contact@eurealimmo.com.",
        align="justify",
    )

    add_heading(doc, "ARTICLE 8 — LITIGES", level=1)
    add_body(doc,
        "En cas de litige, recours amiable préalable obligatoire ; "
        "médiateur de la consommation accessible gratuitement ; à "
        "défaut, juridictions compétentes du domicile du Client.",
        align="justify",
    )

    common_signature_block(doc)
    common_footer(doc)

    return doc


# ============================================================
# Template 3 : MANDAT DE MISE EN LOCATION (recherche locataire)
# ============================================================

def build_mandat_mise_en_location():
    doc = Document()
    set_margins(doc)

    common_header(doc, "DE MISE EN LOCATION (recherche d'un locataire)")
    common_parties(doc)
    common_etat_civil_complement(doc)

    add_heading(doc, "ARTICLE 1 — OBJET DU MANDAT", level=1)
    add_body(doc,
        "Le Client (BAILLEUR) donne mandat à l'Agence, qui l'accepte, "
        "à l'effet de rechercher un locataire pour le bien suivant, "
        "dont le Client déclare être pleinement propriétaire et avoir "
        "le pouvoir de mettre en location :",
        align="justify",
    )
    add_bullet(doc, "Nature du bien : {bien_type}")
    add_bullet(doc, "Adresse du bien : {bien_adresse}")
    add_bullet(doc, "Surface habitable : {bien_surface} m²")
    add_bullet(doc, "Usage : {bien_usage}")
    add_bullet(doc, "Meublé : {bien_meuble}")

    add_heading(doc, "ARTICLE 2 — LOYER ET CHARGES", level=1)
    add_body(doc,
        "Loyer mensuel hors charges : {loyer_hc} EUR ; provisions "
        "sur charges : {charges_mensuelles} EUR ; dépôt de garantie : "
        "{depot_garantie} EUR.",
        align="justify",
    )

    add_heading(doc, "ARTICLE 3 — PÉRIMÈTRE DU MANDAT", level=1)
    add_body(doc,
        "Le présent mandat porte exclusivement sur l'ENTREMISE "
        "INITIALE en vue de la recherche d'un candidat locataire, "
        "et s'achève à la SIGNATURE DU BAIL. Il N'INCLUT PAS la "
        "gestion locative subséquente (encaissement de loyers, "
        "charges, états des lieux post-bail, traitement des "
        "congés), laquelle relève de la carte professionnelle G "
        "dont l'Agence N'EST PAS TITULAIRE.",
        align="justify", italic=True,
    )

    add_heading(doc, "ARTICLE 4 — HONORAIRES DE L'AGENCE", level=1)
    add_body(doc,
        "Les honoraires de l'Agence sont fixés à {commission_mois} "
        "mois de loyer hors charges, soit la somme de "
        "{commission_eur} EUR TTC, conformément à la loi ALUR n° "
        "2014-366 et à son décret d'application n° 2014-890.",
        align="justify", bold=True,
    )
    add_body(doc,
        "Ces honoraires sont répartis entre BAILLEUR ET LOCATAIRE "
        "selon les modalités légales :",
        align="justify",
    )
    add_bullet(doc, "Charge du Bailleur : {commission_bailleur_eur} EUR")
    add_bullet(doc, "Charge du Locataire : {commission_locataire_eur} EUR (plafonnée loi ALUR)")

    common_durée_dénonciation(doc)

    add_heading(doc, "ARTICLE 6 — OBLIGATIONS DU CLIENT (BAILLEUR)", level=1)
    add_bullet(doc, "Fournir les diagnostics obligatoires (DPE, plomb, gaz, électricité, ERP) ;")
    add_bullet(doc, "Garantir le bien en bon état de location, conforme aux normes décence ;")
    add_bullet(doc, "Communiquer titre de propriété, RC propriétaire, justificatifs de paiement de charges ;")
    add_bullet(doc, "Signer le bail dans les meilleurs délais avec le candidat sélectionné conjointement.")

    common_renonciation_l271(doc, signature_distance=True)

    add_heading(doc, "ARTICLE 8 — PROTECTION DES DONNÉES", level=1)
    add_body(doc,
        "Les données du Bailleur et des candidats locataires sont "
        "traitées conformément au RGPD.",
        align="justify",
    )

    add_heading(doc, "ARTICLE 9 — LITIGES", level=1)
    add_body(doc,
        "Recours amiable préalable + médiateur consommation + juridictions compétentes.",
        align="justify",
    )

    common_signature_block(doc)
    common_footer(doc)

    return doc


# ============================================================
# Template 4 : MANDAT DE RECHERCHE DE BIEN À LOUER (locataire/preneur)
# ============================================================

def build_mandat_recherche_bien_locatif():
    doc = Document()
    set_margins(doc)

    common_header(doc, "DE RECHERCHE D'UN BIEN À LOUER")
    common_parties(doc)
    common_etat_civil_complement(doc)

    add_heading(doc, "ARTICLE 1 — OBJET DU MANDAT", level=1)
    add_body(doc,
        "Le Client (CANDIDAT LOCATAIRE / PRENEUR) donne mandat à "
        "l'Agence, qui l'accepte, à l'effet de rechercher pour son "
        "compte un bien immobilier à louer répondant aux critères "
        "suivants :",
        align="justify",
    )
    add_bullet(doc, "Type de bien : {bien_type}")
    add_bullet(doc, "Usage : {bien_usage}")
    add_bullet(doc, "Zone géographique : {bien_zone_geo}")
    add_bullet(doc, "Surface min / max : {bien_surface_min} – {bien_surface_max} m²")
    add_bullet(doc, "Date d'entrée souhaitée : {date_entree_souhaitee}")
    add_bullet(doc, "Critères complémentaires : {bien_criteres_complementaires}")

    add_heading(doc, "ARTICLE 2 — BUDGET LOYER MAXIMUM", level=1)
    add_body(doc,
        "Le Client fixe son budget loyer mensuel maximum charges "
        "incluses à : {loyer_max} EUR.",
        align="justify", bold=True,
    )

    add_heading(doc, "ARTICLE 3 — PÉRIMÈTRE DU MANDAT", level=1)
    add_body(doc,
        "Le présent mandat porte exclusivement sur l'entremise "
        "initiale en vue de la conclusion d'un bail, et s'achève "
        "à la signature dudit bail. La gestion locative ultérieure "
        "ne relève pas du présent mandat (carte G non détenue par "
        "l'Agence).",
        align="justify", italic=True,
    )

    add_heading(doc, "ARTICLE 4 — HONORAIRES DE L'AGENCE", level=1)
    add_body(doc,
        "Les honoraires de l'Agence sont fixés à {commission_mois} "
        "mois de loyer hors charges, soit un montant maximum "
        "estimatif de {commission_eur} EUR TTC. Ces honoraires sont "
        "à la charge du CLIENT (preneur), conformément à la loi "
        "ALUR pour la location de bien à usage d'habitation, ou "
        "selon librement négocié pour les baux commerciaux et "
        "professionnels.",
        align="justify", bold=True,
    )

    common_durée_dénonciation(doc)

    add_heading(doc, "ARTICLE 6 — OBLIGATIONS DU CLIENT", level=1)
    add_bullet(doc, "Fournir les justificatifs nécessaires à la constitution d'un dossier locataire ;")
    add_bullet(doc, "Avertir l'Agence de tout changement de critères ;")
    add_bullet(doc, "S'abstenir de contacter directement un bailleur dont les coordonnées ont été révélées par l'Agence.")

    common_renonciation_l271(doc, signature_distance=True)

    add_heading(doc, "ARTICLE 8 — PROTECTION DES DONNÉES", level=1)
    add_body(doc,
        "Conformément au RGPD, les données du Client sont traitées "
        "pour l'exécution du présent mandat.",
        align="justify",
    )

    add_heading(doc, "ARTICLE 9 — LITIGES", level=1)
    add_body(doc, "Médiation préalable + juridictions compétentes du domicile du Client.",
             align="justify")

    common_signature_block(doc)
    common_footer(doc)

    return doc


# ============================================================
# Main
# ============================================================

def main():
    templates = [
        ("mandat-vente.template.docx", build_mandat_vente),
        ("mandat-recherche-acquereur.template.docx", build_mandat_recherche_acquereur),
        ("mandat-mise-en-location.template.docx", build_mandat_mise_en_location),
        ("mandat-recherche-bien-locatif.template.docx", build_mandat_recherche_bien_locatif),
    ]

    for filename, builder in templates:
        path = OUTPUT_DIR / filename
        doc = builder()
        doc.save(str(path))
        print(f"[OK] {path}")

    print(f"\n{len(templates)} templates Hoguet generes dans {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
