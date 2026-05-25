"""
Génère l'avenant pilote DATAMERRY × Collabimmo en .docx.

Usage:
    python scripts/generate-avenant-collabimmo.py

Sortie:
    public/legal/avenant-pilote-collabimmo.docx

Le document est en français, format A4, police Arial, prêt à être signé
électroniquement (DocuSign / Yousign) ou imprimé/signé/scanné.
"""

from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


PRIMARY = RGBColor(0x1F, 0x3A, 0x8A)  # bleu DATAMERRY
GREY = RGBColor(0x64, 0x74, 0x8B)
TODAY = date.today().strftime("%d/%m/%Y")


def set_cell_border(cell, color="CCCCCC", sz="4"):
    """Bordures fines grises pour les cellules."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_heading(doc, text, level=1):
    """Titres bleu DATAMERRY, Arial bold."""
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.bold = True
    run.font.color.rgb = PRIMARY
    run.font.size = Pt(18 if level == 0 else 14 if level == 1 else 12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_para(doc, text, bold=False, italic=False, size=11, align=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    """Bullet point Arial 11pt avec retrait."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.75)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)
    if not p.runs:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(11)
    else:
        p.runs[0].text = text
    return p


def add_two_party_block(doc):
    """Tableau 1×2 'entre les soussignées' format contrat FR."""
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8.5)
    table.columns[1].width = Cm(8.5)
    cells = table.rows[0].cells

    # Eurealimmo
    cells[0].width = Cm(8.5)
    c0 = cells[0].paragraphs[0]
    c0.add_run("EUREALIMMO SARL\n").bold = True
    c0.add_run(
        "Société à responsabilité limitée à associé unique\n"
        "Capital : 100 €\n"
        "Siège : 60 rue François 1er, 75008 Paris\n"
        "RCS Paris : [N° KBIS à compléter]\n"
        "Représentée par M. Samuel BRUNO, Gérant\n"
        "\n"
        "Ci-après dénommée « l'Éditeur »\n"
        "ou « DATAMERRY »"
    )
    for run in c0.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)

    # Collabimmo
    cells[1].width = Cm(8.5)
    c1 = cells[1].paragraphs[0]
    c1.add_run("COLLABIMMO\n").bold = True
    c1.add_run(
        "[Forme juridique à compléter]\n"
        "[Capital à compléter]\n"
        "[Siège social à compléter]\n"
        "[RCS à compléter]\n"
        "Représentée par Mme Diara CAMARA, [fonction]\n"
        "\n"
        "Ci-après dénommée « le Client »\n"
        "ou « Collabimmo »"
    )
    for run in c1.runs:
        run.font.name = "Arial"
        run.font.size = Pt(10)

    for cell in cells:
        set_cell_border(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    return table


def add_signatures_block(doc):
    """Bloc signatures sur 2 colonnes."""
    table = doc.add_table(rows=2, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8.5)
    table.columns[1].width = Cm(8.5)

    # Ligne 1 : titres
    titles = table.rows[0].cells
    for i, label in enumerate(["Pour l'Éditeur — EUREALIMMO SARL", "Pour le Client — COLLABIMMO"]):
        p = titles[i].paragraphs[0]
        run = p.add_run(label)
        run.font.name = "Arial"
        run.font.bold = True
        run.font.size = Pt(10)
        set_cell_border(titles[i])

    # Ligne 2 : zone signature
    sigs = table.rows[1].cells
    for i, name in enumerate(["M. Samuel BRUNO\nGérant", "Mme Diara CAMARA\n[Fonction]"]):
        p = sigs[i].paragraphs[0]
        p.add_run(f"Fait à _______________ le {TODAY}\n\n").font.name = "Arial"
        p.add_run("Signature (précédée de la mention « Lu et approuvé ») :\n\n\n\n\n").font.name = "Arial"
        run = p.add_run(name)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.bold = True
        for r in p.runs:
            r.font.size = Pt(10) if r.font.size is None else r.font.size
            r.font.name = "Arial"
        sigs[i].vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_border(sigs[i])

    return table


def main():
    out = Path(__file__).resolve().parent.parent / "public" / "legal" / "avenant-pilote-collabimmo.docx"
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Marges A4
    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Style par défaut
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ─── Titre ───
    add_heading(doc, "AVENANT N° 1 — OFFRE PILOTE DATAMERRY", level=0)
    add_para(
        doc,
        "Au contrat-cadre de prestation de service numérique DATAMERRY®",
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=GREY,
        size=11,
    )
    add_para(doc, f"En date du {TODAY}", align=WD_ALIGN_PARAGRAPH.CENTER, color=GREY, size=10)

    # ─── Entre les soussignées ───
    add_para(doc, "Entre les soussignées :", bold=True, size=11)
    add_two_party_block(doc)
    add_para(doc, "Il a été préalablement exposé ce qui suit :", bold=True, size=11)

    # ─── Préambule ───
    add_heading(doc, "PRÉAMBULE", level=1)
    add_para(
        doc,
        "L'Éditeur exploite, sous la marque DATAMERRY®, une plateforme d'estimation immobilière "
        "et d'analyse de marché à destination des professionnels titulaires de la carte T "
        "(loi Hoguet) et de leurs mandataires. La plateforme intègre les sources publiques DVF, "
        "OLAP, ANIL, INSEE, Cadastre et OpenStreetMap pour produire des indicateurs d'évaluation, "
        "de rendement locatif, de plafonds fiscaux et un rapport propriété multi-datasets.",
    )
    add_para(
        doc,
        "Le Client, agence immobilière partenaire, souhaite tester le Service en avant-première "
        "et bénéficier en contrepartie d'un tarif avantageux gelé pour la durée de son utilisation, "
        "en qualité de cabinet pilote n° 1.",
    )
    add_para(
        doc,
        "Les Parties se sont rapprochées et ont convenu des dispositions qui suivent. Le présent "
        "avenant complète les Conditions Générales de Vente et d'Utilisation (CGV) du Service "
        "DATAMERRY® en vigueur, qui demeurent applicables pour tout point non expressément modifié "
        "par le présent document.",
        italic=True,
    )

    # ─── Article 1 — Objet ───
    add_heading(doc, "ARTICLE 1 — OBJET DE L'AVENANT", level=1)
    add_para(
        doc,
        "Le présent avenant a pour objet de définir les conditions spécifiques applicables au "
        "Client en sa qualité de Cabinet Pilote n° 1 du Service DATAMERRY®, notamment :",
    )
    add_bullet(doc, "le tarif gelé applicable à perpétuité au Service v1 (Article 2) ;")
    add_bullet(doc, "le quota d'utilisation inclus et les modalités de dépassement (Article 3) ;")
    add_bullet(doc, "l'engagement de référence consenti par le Client (Article 4) ;")
    add_bullet(doc, "le périmètre exclu du tarif gelé (Article 5).")

    # ─── Article 2 — Tarif gelé ───
    add_heading(doc, "ARTICLE 2 — TARIF GELÉ", level=1)
    add_para(
        doc,
        "L'Éditeur consent au Client un tarif gelé de TRENTE-NEUF EUROS toutes taxes comprises "
        "(39,00 € TTC) par mois pour l'ensemble du Service v1 défini à l'Article 5.1 ci-après.",
        bold=True,
    )
    add_para(
        doc,
        "Ce tarif est gelé pour toute la durée pendant laquelle le Client conserve son "
        "abonnement actif et continu, sans interruption supérieure à 90 jours consécutifs. Aucune "
        "augmentation tarifaire ne pourra être appliquée au Client tant qu'il bénéficie du "
        "présent gel, et ce indépendamment des évolutions tarifaires que l'Éditeur pourrait "
        "appliquer à d'autres clients.",
    )
    add_para(
        doc,
        "Le premier mois d'abonnement est offert au Client, conformément à l'offre commerciale "
        "standard de l'Éditeur.",
    )

    # ─── Article 3 — Quota et dépassement ───
    add_heading(doc, "ARTICLE 3 — QUOTA INCLUS ET MODALITÉS DE DÉPASSEMENT", level=1)
    add_para(
        doc,
        "Le tarif gelé visé à l'Article 2 inclut un quota de CINQUANTE MILLE (50 000) requêtes "
        "API par mois calendaire, applicable indistinctement à l'ensemble des endpoints du "
        "Service v1.",
    )
    add_para(doc, "En cas de dépassement de ce quota, l'Éditeur procédera, au choix du Client :")
    add_bullet(doc, "soit à la suspension automatique des requêtes excédentaires jusqu'au mois suivant ;")
    add_bullet(
        doc,
        "soit, sur demande écrite préalable du Client, à la facturation des requêtes "
        "excédentaires au tarif standard en vigueur (1,00 € TTC par tranche de 1 000 requêtes), "
        "sans remise.",
    )
    add_para(
        doc,
        "Il est expressément convenu que le quota inclus ne constitue pas un droit à un usage "
        "illimité gratuit. Le Client s'interdit toute revente ou redistribution à des tiers de "
        "son accès, ainsi que toute extraction massive systématique des données.",
        italic=True,
    )

    # ─── Article 4 — Engagement de référence ───
    add_heading(doc, "ARTICLE 4 — ENGAGEMENT DE RÉFÉRENCE DU CLIENT", level=1)
    add_para(
        doc,
        "En contrepartie du tarif gelé consenti au présent avenant, le Client accepte expressément :",
    )
    add_bullet(
        doc,
        "l'affichage de sa dénomination commerciale « Collabimmo » et de son logo sur le site "
        "internet de l'Éditeur (datamerry.com) dans la section « Cabinets de référence » ou "
        "équivalent ;",
    )
    add_bullet(
        doc,
        "la mise à disposition d'un témoignage écrit court (1 à 3 phrases) attestant de "
        "l'utilisation et de la satisfaction du Service, accompagné d'une photographie ou d'un "
        "portrait de Mme Diara CAMARA si elle l'autorise ;",
    )
    add_bullet(
        doc,
        "le droit pour l'Éditeur d'utiliser cette référence dans ses supports commerciaux "
        "(plaquette, présentation à des prospects, communiqués de presse) sous réserve d'un "
        "usage loyal et conforme à l'image du Client.",
    )
    add_para(
        doc,
        "Le Client peut demander à tout moment, par simple email à contact@datamerry.com, le "
        "retrait de sa référence. Dans cette hypothèse, le tarif standard de l'Éditeur en vigueur "
        "à la date de la demande de retrait s'applique automatiquement à compter du mois suivant, "
        "le gel tarifaire de l'Article 2 cessant alors de produire ses effets.",
    )

    # ─── Article 5 — Périmètre du gel ───
    add_heading(doc, "ARTICLE 5 — PÉRIMÈTRE DU GEL", level=1)
    add_heading(doc, "5.1 Service v1 couvert par le gel", level=2)
    add_para(
        doc,
        "Le tarif gelé visé à l'Article 2 couvre exclusivement le périmètre fonctionnel suivant, "
        "ci-après dénommé « Service v1 » :",
    )
    add_bullet(doc, "endpoint /api/address/search (résolution d'adresses par IA) ;")
    add_bullet(doc, "endpoint /api/estimate (estimation par cluster HDBSCAN) ;")
    add_bullet(doc, "endpoint /api/rendement (rendement locatif OLAP/ANIL) ;")
    add_bullet(doc, "endpoint /api/plafonds-fiscaux (Jeanbrun, LLI, Loc'Avantages, Denormandie) ;")
    add_bullet(doc, "endpoint /api/rental-strategies (8 scénarios locatifs comparés) ;")
    add_bullet(doc, "endpoint /api/property-report (rapport agrégé : estimation + rendement + plafonds + écoles + transports + services proximité + streetview) ;")
    add_bullet(doc, "endpoint /api/property-report/pdf (export PDF brandé) ;")
    add_bullet(doc, "endpoint /api/widget/render (rendu HTML stylé) et widget JavaScript embarquable.")

    add_heading(doc, "5.2 Fonctionnalités futures exclues du gel", level=2)
    add_para(
        doc,
        "Toute nouvelle fonctionnalité majeure ajoutée au catalogue de l'Éditeur après la "
        "signature du présent avenant fera l'objet d'une tarification distincte au tarif marché "
        "en vigueur à la date de sa mise à disposition. Sont notamment visées, sans que cette "
        "liste soit limitative :",
    )
    add_bullet(doc, "le module « Combo carte T » (hébergement de mandataires sous la carte T de l'Éditeur, génération de mandats et compromis, reversement de commissions) ;")
    add_bullet(doc, "le chatbot d'analyse LLM (Phase 3 — post-formation IBM) ;")
    add_bullet(doc, "les datasets premium additionnels (cadastre détaillé, permis de construire, INSEE socio-démo IRIS, etc.) ;")
    add_bullet(doc, "tout nouveau dispositif fiscal ou réglementaire intégré ultérieurement.")
    add_para(
        doc,
        "Le Client pourra librement souscrire ou décliner ces fonctionnalités futures sans que "
        "cela n'affecte le gel tarifaire applicable au Service v1.",
        italic=True,
    )

    # ─── Article 6 — Durée et résiliation ───
    add_heading(doc, "ARTICLE 6 — DURÉE ET RÉSILIATION", level=1)
    add_para(
        doc,
        "Le présent avenant prend effet à la date de sa signature par les deux Parties et "
        "demeure en vigueur tant que le Client conserve son abonnement actif au Service v1, "
        "sans interruption supérieure à 90 jours consécutifs.",
    )
    add_para(
        doc,
        "En cas d'interruption de l'abonnement supérieure à 90 jours consécutifs, le présent "
        "avenant cesse automatiquement de produire ses effets ; toute nouvelle souscription du "
        "Client interviendra alors aux conditions standard en vigueur au jour de la nouvelle "
        "souscription.",
    )
    add_para(
        doc,
        "Le présent avenant peut également être résilié par l'une ou l'autre des Parties, sans "
        "préavis et sans pénalité, en cas de manquement grave de l'autre Partie à ses "
        "obligations, après mise en demeure restée infructueuse pendant 30 jours.",
    )

    # ─── Article 7 — Bascule juridique ───
    add_heading(doc, "ARTICLE 7 — BASCULE ÉDITORIALE EUREALIMMO → DATAMERRY SAS", level=1)
    add_para(
        doc,
        "Les Parties reconnaissent que l'entité juridique DATAMERRY SAS est en cours "
        "d'immatriculation et que, pendant la phase transitoire, le Service est exploité et "
        "facturé par EUREALIMMO SARL en sa qualité d'éditeur. Dès l'immatriculation effective "
        "de DATAMERRY SAS, l'ensemble des droits et obligations du présent avenant sera "
        "transféré de plein droit à DATAMERRY SAS par voie de cession de fonds de commerce "
        "numérique, sans qu'aucune action ne soit requise du Client.",
    )
    add_para(
        doc,
        "Le Client en sera informé par email au moins trente (30) jours avant la bascule. "
        "Les conditions financières, le périmètre et l'ensemble des stipulations du présent "
        "avenant demeureront strictement identiques.",
        italic=True,
    )

    # ─── Article 8 — Droit applicable ───
    add_heading(doc, "ARTICLE 8 — DROIT APPLICABLE ET JURIDICTION", level=1)
    add_para(
        doc,
        "Le présent avenant est régi par le droit français. Tout litige relatif à sa formation, "
        "son exécution ou son interprétation sera soumis à la compétence exclusive des "
        "tribunaux de Paris, après tentative préalable de règlement amiable entre les Parties.",
    )

    # ─── Signatures ───
    add_heading(doc, "SIGNATURES", level=1)
    add_para(
        doc,
        "Fait en deux exemplaires originaux, un pour chaque Partie.",
        italic=True,
    )
    doc.add_paragraph()  # espace
    add_signatures_block(doc)

    # ─── Footer notice ───
    doc.add_paragraph()
    add_para(
        doc,
        "Document v1.0 — généré le " + TODAY + " — Service DATAMERRY® exploité par EUREALIMMO SARL "
        "(en attente de bascule sur DATAMERRY SAS en cours d'immatriculation).",
        italic=True,
        size=8,
        color=GREY,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.save(str(out))
    print(f"OK -> {out}")
    print(f"Taille : {out.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
